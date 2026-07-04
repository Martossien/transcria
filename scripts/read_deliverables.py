#!/usr/bin/env python3
"""Extrait le CONTENU RÉEL des livrables d'un job pour lecture humaine (validation E2E).

Le script de campagne vérifie la STRUCTURE ; celui-ci sort la SUBSTANCE (texte SRT,
résumé et participants du Word, points qualité) pour qu'on la LISE — un livrable peut
être structurellement valide et vide de sens (« Résumé indisponible », charabia).

    python scripts/read_deliverables.py --jobs-dir <dir> --job <id>
"""
from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path


def _docx_text(docx_path: Path) -> str:
    """Texte lisible du DOCX (python-docx : paragraphes + cellules de tableau)."""
    from docx import Document

    doc = Document(str(docx_path))
    lines: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(dict.fromkeys(cells)))
    return "\n".join(lines)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Lire les livrables réels d'un job")
    parser.add_argument("--jobs-dir", required=True)
    parser.add_argument("--job", required=True)
    parser.add_argument("--srt-lines", type=int, default=25, help="nb de lignes SRT à montrer")
    args = parser.parse_args(argv)

    job_dir = Path(args.jobs_dir) / args.job
    if not job_dir.is_dir():
        print(f"job introuvable : {job_dir}", file=sys.stderr)
        return 1

    print("=" * 72)
    print(f"LIVRABLES DU JOB {args.job}")
    print("=" * 72)

    # ── SRT corrigé (le livré) ────────────────────────────────────────────────
    srt = (job_dir / "metadata" / "transcription_corrigee.srt")
    if not srt.exists():
        srt = job_dir / "metadata" / "transcription.srt"
    if srt.exists():
        lines = srt.read_text(encoding="utf-8").splitlines()
        text_lines = [ln for ln in lines if ln and "-->" not in ln and not ln.strip().isdigit()]
        print(f"\n── SRT ({srt.name}) : {len(text_lines)} lignes de texte ──")
        # locuteurs distincts
        speakers = sorted({m.group(1) for ln in text_lines
                           if (m := re.match(r"(SPEAKER_\d+(?:\([^)]*\))?)", ln))})
        print(f"   Locuteurs : {', '.join(speakers) or '(aucun préfixe)'}")
        print("   — Extrait —")
        for ln in text_lines[:args.srt_lines]:
            print(f"   {ln}")
    else:
        print("\n⚠ AUCUN SRT trouvé")

    # ── Résumé (summary.md effectif) ──────────────────────────────────────────
    summary = job_dir / "summary" / "summary.md"
    if summary.exists():
        content = summary.read_text(encoding="utf-8")
        print(f"\n── RÉSUMÉ (summary.md, {len(content)} car.) ──")
        print("\n".join("   " + ln for ln in content.splitlines()[:40]))
    else:
        print("\n⚠ AUCUN summary.md")

    # ── DOCX : texte du rapport ───────────────────────────────────────────────
    docx = next((job_dir / "exports").glob("*.docx"), None) if (job_dir / "exports").is_dir() else None
    if docx and docx.stat().st_size > 200:
        try:
            text = _docx_text(docx)
            print(f"\n── DOCX ({docx.name}, {len(text)} car. de texte) ──")
            print("\n".join("   " + ln for ln in text.splitlines() if ln.strip())[:3000])
        except Exception as exc:  # noqa: BLE001
            print(f"\n⚠ DOCX illisible : {exc}")
    else:
        print("\n⚠ AUCUN DOCX exploitable dans exports/")

    # ── Points qualité ────────────────────────────────────────────────────────
    rp = job_dir / "quality" / "review_points.json"
    if rp.exists():
        import json
        points = json.loads(rp.read_text(encoding="utf-8"))
        print(f"\n── POINTS À VÉRIFIER ({len(points)}) ──")
        for p in points[:12]:
            print(f"   • {p}")

    print()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

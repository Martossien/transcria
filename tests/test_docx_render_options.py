"""Options de rendu DOCX data-driven (``context/render_options.json``) — GPU-free.

Contrat : la LLM (ou l'utilisateur, via l'UI) choisit *quoi* rendre ; le renderer
garantit *comment*. Options v1 : ``theme`` (clé de ``_THEMES``, prime sur le
``meeting_type``) et ``sections`` (booléens : ``participants`` / ``transcript`` /
``quality``). Tout invalide est ignoré silencieusement — le rendu ne casse JAMAIS.
"""
import json
from pathlib import Path

import pytest
from test_docx_report import _seed_job

pytest.importorskip("docx")


def _render(tmp_path, options: dict | None) -> "object":
    from docx import Document

    from transcria.exports.docx_report import generate_docx_report

    job_id = "job-options"
    _seed_job(tmp_path, job_id)
    if options is not None:
        (tmp_path / job_id / "context" / "render_options.json").write_text(
            json.dumps(options), encoding="utf-8"
        )
    out = tmp_path / "rapport.docx"
    generate_docx_report(job_id, str(tmp_path), out)
    return Document(str(out))


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    parts += [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(parts)


class TestThemeOverride:
    def test_theme_overrides_meeting_type(self, tmp_path):
        # meeting_type="Autre" mais theme="CSE" → bandeau CSE appliqué.
        doc = _render(tmp_path, {"theme": "CSE"})
        assert "PROCÈS-VERBAL DU COMITÉ SOCIAL ET ÉCONOMIQUE" in _all_text(doc)

    def test_unknown_theme_ignored(self, tmp_path):
        ref = _all_text(_render(tmp_path, None))
        doc = _all_text(_render(tmp_path, {"theme": "zzz-inexistant"}))
        assert doc == ref  # rendu par défaut, aucune exception


def _section_headings(doc) -> list[str]:
    """Titres de section numérotés (« N.  LIBELLÉ ») — ignore le bandeau de couverture."""
    import re

    return [p.text for p in doc.paragraphs if re.match(r"^\d+\.\s{2}\S", p.text)]


class TestSectionToggles:
    def test_default_has_all_sections(self, tmp_path):
        heads = "\n".join(_section_headings(_render(tmp_path, None)))
        assert "TRANSCRIPTION" in heads
        assert "PARTICIPANTS & LOCUTEURS" in heads
        assert "POINTS À VÉRIFIER" in heads

    def test_transcript_off(self, tmp_path):
        heads = "\n".join(_section_headings(_render(tmp_path, {"sections": {"transcript": False}})))
        assert "TRANSCRIPTION" not in heads
        assert "PARTICIPANTS & LOCUTEURS" in heads  # les autres restent

    def test_participants_and_quality_off(self, tmp_path):
        heads = "\n".join(_section_headings(
            _render(tmp_path, {"sections": {"participants": False, "quality": False}})))
        assert "PARTICIPANTS & LOCUTEURS" not in heads
        assert "POINTS À VÉRIFIER" not in heads
        assert "TRANSCRIPTION" in heads

    def test_numbering_stays_sequential_when_section_skipped(self, tmp_path):
        # participants off → les numéros restent séquentiels (pas de trou 2. → 4.).
        doc = _render(tmp_path, {"sections": {"participants": False}})
        nums = [int(h.split(".")[0]) for h in _section_headings(doc)]
        assert nums == list(range(1, len(nums) + 1)), f"numérotation non séquentielle : {nums}"


class TestRobustness:
    def test_garbage_options_ignored(self, tmp_path):
        ref = _all_text(_render(tmp_path, None))
        doc = _all_text(_render(tmp_path, {"theme": 42, "sections": "junk", "autre": []}))
        assert doc == ref

    def test_corrupt_json_ignored(self, tmp_path):
        job_id = "job-options"
        _seed_job(tmp_path, job_id)
        (tmp_path / job_id / "context" / "render_options.json").write_text("{pas du json", encoding="utf-8")
        from transcria.exports.docx_report import generate_docx_report
        out = tmp_path / "rapport.docx"
        generate_docx_report(job_id, str(tmp_path), out)  # ne lève pas
        assert Path(out).is_file()


class TestSpeakerNamesInSummary:
    """Vécu 2026-07-19 : la synthèse (rédigée par la LLM AVANT la validation des
    locuteurs) parlait de SPEAKER_00 dans le DOCX alors que les noms étaient
    validés — substitution au RENDU, artefacts intacts."""

    def test_substitute_speaker_names_regles(self):
        from transcria.workflow.speaker_projection import substitute_speaker_names

        m = {"mapping": {"SPEAKER_00": {"name": "Alice Durand"},
                         "SPEAKER_01": {"name": ""},            # non renseigné → conservé
                         "SPEAKER_02": {"name": "SPEAKER_02"}}}  # placeholder → conservé
        out = substitute_speaker_names(
            "SPEAKER_00 répond à SPEAKER_01 puis SPEAKER_02. Voir SPEAKER_00.", m)
        assert out == "Alice Durand répond à SPEAKER_01 puis SPEAKER_02. Voir Alice Durand."
        # Jamais de correspondance partielle ni de texte sans mapping modifié.
        assert substitute_speaker_names("SPEAKER_003 intact", m) == "SPEAKER_003 intact"
        assert substitute_speaker_names("texte", None) == "texte"

    def test_docx_synthese_et_actions_avec_noms(self, tmp_path):
        import re

        from docx import Document

        from transcria.exports.docx_report import generate_docx_report
        from transcria.jobs.filesystem import JobFilesystem

        fs = JobFilesystem(str(tmp_path), "j1")
        fs.save_json("context/meeting_context.json", {
            "summary_llm": "SPEAKER_00 présente le budget. SPEAKER_01 valide.",
            "structured_data": {"actions": ["SPEAKER_01 : envoyer le compte rendu"]},
        })
        fs.save_json("context/participants.json", [])
        fs.save_json("speakers/speaker_mapping.json",
                     {"mapping": {"SPEAKER_00": {"name": "Alice Durand"},
                                  "SPEAKER_01": {"name": "Bob Martin"}}})
        fs.save_text("metadata/transcription.srt",
                     "1\n00:00:01,000 --> 00:00:02,000\nSPEAKER_00(Alice Durand): Bonjour\n")

        out = tmp_path / "rapport.docx"
        generate_docx_report("j1", str(tmp_path), out)
        text = "\n".join(p.text for p in Document(str(out)).paragraphs)
        assert "Alice Durand présente le budget" in text
        assert "Bob Martin : envoyer le compte rendu" in text
        assert not re.search(r"SPEAKER_\d{2}(?!\()", text)

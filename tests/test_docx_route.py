"""
Tests E2E du endpoint GET /api/jobs/<id>/download/docx.

Couvre :
  - Génération et téléchargement du rapport DOCX via HTTP
  - Authentification requise (401/302 si non connecté)
  - 404 sur job inexistant
  - Contenu DOCX valide (magic bytes, sections, participants)
  - Page résultat contient le bouton de téléchargement
  - ZIP PackageBuilder intègre le DOCX
"""
import zipfile

import pytest

# ── Fixtures communes ─────────────────────────────────────────────────────────

def _seed_job_files(jobs_dir: str, job_id: str) -> None:
    """Remplit le répertoire d'un job avec les fichiers nécessaires au DOCX."""
    from transcria.jobs.filesystem import JobFilesystem

    fs = JobFilesystem(jobs_dir, job_id)

    fs.save_json("context/meeting_context.json", {
        "title": "Réunion de test DOCX",
        "meeting_type": "Réunion interne",
        "date": "2026-05-29",
        "service": "Département IT",
        "language": "fr",
        "topic": "Validation du rapport Word",
        "objective": "Vérifier que le rapport DOCX est généré correctement.",
        "notes": "Test automatique.",
        "summary": "La réunion s'est bien déroulée. L'équipe a validé la feature DOCX.",
        "sensitivity": "normal",
    })
    fs.save_json("context/participants.json", [
        {
            "id": "p1", "name": "Alice Dupont", "function": "Chef de projet",
            "service": "IT", "role": "Animatrice", "is_animator": True,
            "expected": True, "comment": "",
        },
        {
            "id": "p2", "name": "Bob Martin", "function": "Développeur",
            "service": "IT", "role": "Contributeur", "is_animator": False,
            "expected": True, "comment": "",
        },
    ])
    fs.save_json("speakers/speaker_stats.json", {
        "speakers": [
            {
                "speaker_id": "SPEAKER_00", "label": "SPEAKER_00",
                "mapped_to": "p1", "mapped_name": "Alice Dupont",
                "speaking_time_seconds": 40.0, "turn_count": 8,
                "validation": "user_validated", "gender": "female",
            },
            {
                "speaker_id": "SPEAKER_01", "label": "SPEAKER_01",
                "mapped_to": "p2", "mapped_name": "Bob Martin",
                "speaking_time_seconds": 20.0, "turn_count": 4,
                "validation": "user_validated", "gender": "male",
            },
        ]
    })
    fs.save_text(
        "metadata/transcription_corrigee.srt",
        "1\n00:00:01,000 --> 00:00:03,500\n"
        "SPEAKER_00(Alice Dupont): Bonjour à tous, on commence la réunion.\n\n"
        "2\n00:00:04,000 --> 00:00:06,000\n"
        "SPEAKER_01(Bob Martin): Oui, bonne idée.\n\n"
        "3\n00:00:07,000 --> 00:00:10,000\n"
        "SPEAKER_00(Alice Dupont): Le rapport Word est maintenant généré automatiquement.\n\n",
    )
    fs.save_json("quality/quality_report.json", {
        "quality_score": 90,
        "total_checks": 10,
        "warnings": 0,
        "checks": [],
        "review_points": [],
    })


def _make_job_id(admin_client) -> str | None:
    r = admin_client.post("/jobs/new", data={"title": "Test DOCX E2E"}, follow_redirects=True)
    path = r.request.path
    return path.split("/")[2] if "/jobs/" in path else None


@pytest.fixture
def job_with_docx_data(admin_client, app):
    """Crée un job et peuple ses fichiers, retourne son job_id."""
    job_id = _make_job_id(admin_client)
    assert job_id, "Impossible de créer le job de test"

    from transcria.config import get_config
    cfg = get_config()
    with app.app_context():
        _seed_job_files(cfg["storage"]["jobs_dir"], job_id)

    return job_id


# ── Tests authentification ────────────────────────────────────────────────────

class TestDocxAuth:
    def test_download_docx_requires_login(self, client):
        r = client.get("/api/jobs/any-id/download/docx")
        assert r.status_code in (302, 401)

    def test_download_docx_unknown_job_returns_404(self, admin_client):
        r = admin_client.get("/api/jobs/nonexistent-uuid-docx/download/docx")
        assert r.status_code == 404

    def test_download_docx_verbatim_pour_profil_srt(self, admin_client, app, job_with_docx_data):
        # Profil SRT (docx_level == none) : depuis 0.3.8, DOCX VERBATIM généré à la
        # demande (PISTES_AMELIORATION §5.1) — capacité additive, les livrables du
        # profil (ZIP minimal, aucune passe LLM) ne changent pas.
        from transcria.jobs.store import JobStore
        with app.app_context():
            JobStore.update_extra_data(
                job_with_docx_data,
                lambda d: {**d, "execution": {"processing_profile_id": "srt_express"}},
            )
        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        assert r.status_code == 200
        assert r.data[:4] == b"PK\x03\x04"  # DOCX = zip valide

    def test_download_docx_verbatim_degrade_sans_artefacts_llm(self, admin_client, app):
        # Job SRT MINIMAL (aucun artefact LLM/diarisation : ni synthèse, ni
        # participants, ni stats locuteurs) : le générateur dégrade proprement —
        # jamais de 500, un document lisible avec le verbatim.
        from transcria.jobs.filesystem import JobFilesystem
        from transcria.jobs.models import JobState
        from transcria.jobs.store import JobStore

        with app.app_context():
            from transcria.auth.store import UserStore
            from transcria.config import get_config
            jobs_dir = get_config()["storage"]["jobs_dir"]
            owner = UserStore.get_by_username("admin")
            job = JobStore.create_job(owner.id, "SRT express nu")
            JobStore.update_state(job.id, JobState.COMPLETED)
            JobStore.update_extra_data(
                job.id,
                lambda d: {**d, "execution": {"processing_profile_id": "srt_express"}},
            )
            fs = JobFilesystem(jobs_dir, job.id)
            fs.save_text("metadata/transcription.srt",
                         "1\n00:00:00,000 --> 00:00:02,000\nBonjour à tous.\n\n"
                         "2\n00:00:02,500 --> 00:00:04,000\nOn commence.\n")
            job_id = job.id

        r = admin_client.get(f"/api/jobs/{job_id}/download/docx")
        assert r.status_code == 200
        assert r.data[:4] == b"PK\x03\x04"

    def test_download_docx_200_pour_profil_word(self, admin_client, app, job_with_docx_data):
        # Profil Word (docx_level != none) : le DOCX reste un livrable.
        from transcria.jobs.store import JobStore
        with app.app_context():
            JobStore.update_extra_data(
                job_with_docx_data,
                lambda d: {**d, "execution": {"processing_profile_id": "dossier_qualite"}},
            )
        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        assert r.status_code == 200


# ── Tests génération et contenu ───────────────────────────────────────────────

class TestDocxDownload:
    def test_returns_200(self, admin_client, job_with_docx_data):
        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        assert r.status_code == 200

    def test_content_type_docx(self, admin_client, job_with_docx_data):
        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        assert "openxmlformats" in r.content_type or "docx" in r.content_type

    def test_content_disposition_attachment(self, admin_client, job_with_docx_data):
        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        cd = r.headers.get("Content-Disposition", "")
        assert "attachment" in cd
        assert ".docx" in cd

    def test_response_is_valid_docx(self, admin_client, job_with_docx_data):
        """Un DOCX est un ZIP — vérifie les magic bytes PK."""
        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        assert r.data[:4] == b"PK\x03\x04"

    def test_docx_taille_minimale(self, admin_client, job_with_docx_data):
        """Le fichier doit peser au moins 5 Ko — sinon il est vide ou corrompu."""
        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        assert len(r.data) > 5_000

    def test_docx_contient_sections_attendues(self, admin_client, job_with_docx_data, tmp_path):
        pytest.importorskip("docx")
        import io

        from docx import Document

        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        doc = Document(io.BytesIO(r.data))

        all_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
        full = all_text + " " + table_text

        assert "CONTEXTE"       in full.upper()
        assert "PARTICIPANTS"   in full.upper()
        assert "TRANSCRIPTION"  in full.upper()

    def test_docx_contient_noms_participants(self, admin_client, job_with_docx_data):
        pytest.importorskip("docx")
        import io

        from docx import Document

        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        doc = Document(io.BytesIO(r.data))
        table_text = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)

        assert "Alice Dupont" in table_text
        assert "Bob Martin"   in table_text

    def test_docx_contient_transcription(self, admin_client, job_with_docx_data):
        pytest.importorskip("docx")
        import io

        from docx import Document

        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        doc = Document(io.BytesIO(r.data))
        table_text = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)

        assert "Alice Dupont"   in table_text
        assert "Bob Martin"     in table_text
        assert "00:00:01"       in table_text

    def test_docx_pourcentages_temps_parole(self, admin_client, job_with_docx_data):
        """40s + 20s → Alice 67%, Bob 33% — doit apparaître dans le tableau."""
        pytest.importorskip("docx")
        import io
        import re

        from docx import Document

        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        doc = Document(io.BytesIO(r.data))
        table_text = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)

        assert re.search(r"6[67]%", table_text)
        assert re.search(r"3[23]%", table_text)

    def test_docx_sans_section_qualite_si_score_ok(self, admin_client, job_with_docx_data):
        """Score 90, pas de flags → le titre de section '4.' absent du document."""
        pytest.importorskip("docx")
        import io

        from docx import Document

        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        doc = Document(io.BytesIO(r.data))
        # La section qualité est ajoutée avec le numéro "4." — absente si aucun flag
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "4." not in all_text

    def test_docx_section_qualite_si_coverage_faible(self, admin_client, app):
        """Coverage 70% → section 'Points à vérifier' présente."""
        pytest.importorskip("docx")
        import io

        from docx import Document

        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem

        # Nouveau job avec coverage faible
        job_id = _make_job_id(admin_client)
        assert job_id

        with app.app_context():
            cfg = get_config()
            fs = JobFilesystem(cfg["storage"]["jobs_dir"], job_id)
            fs.save_json("context/meeting_context.json", {"title": "Test", "language": "fr"})
            fs.save_json("context/participants.json", [])
            fs.save_json("speakers/speaker_stats.json", {"speakers": []})
            fs.save_text("metadata/transcription_corrigee.srt", "")
            fs.save_json("quality/quality_report.json", {
                "quality_score": 65,
                "total_checks": 5,
                "warnings": 1,
                "checks": [
                    {"type": "low_coverage", "ratio": 0.70, "severity": "error"}
                ],
                "review_points": ["Couverture faible"],
            })

        r2 = admin_client.get(f"/api/jobs/{job_id}/download/docx")
        assert r2.status_code == 200
        doc = Document(io.BytesIO(r2.data))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
        full = all_text + " " + table_text
        assert "VÉRIFIER" in full.upper() or "70" in full

    def test_docx_cached_sur_disque(self, admin_client, app, job_with_docx_data):
        """Deux appels successifs → le fichier est mis en cache dans exports/."""

        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem

        # Premier appel
        r1 = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        assert r1.status_code == 200

        with app.app_context():
            cfg = get_config()
            fs = JobFilesystem(cfg["storage"]["jobs_dir"], job_with_docx_data)
            # Vérifier qu'un .docx existe dans exports/
            docx_files = list((fs.job_dir / "exports").glob("*.docx"))
            assert docx_files, "Le DOCX doit être sauvegardé dans exports/"

        # Deuxième appel — doit aussi fonctionner (régénération propre)
        r2 = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/docx")
        assert r2.status_code == 200


# ── Tests page résultat ───────────────────────────────────────────────────────

class TestJobResultPage:
    def test_bouton_docx_present_dans_page_resultat(self, admin_client, app, job_with_docx_data):
        # /result n'est servie que pour un job COMPLETED (garde R2 de la revue macro)
        with app.app_context():
            from transcria.jobs.models import JobState
            from transcria.jobs.store import JobStore
            JobStore.update_state(job_with_docx_data, JobState.COMPLETED)
        r = admin_client.get(f"/jobs/{job_with_docx_data}/result")
        assert r.status_code == 200
        body = r.data.decode("utf-8")
        assert "download/docx" in body
        assert "Rapport Word" in body or "docx" in body.lower()


# ── Tests ZIP intègre le DOCX ─────────────────────────────────────────────────

class TestZipIncludesDocx:
    def test_package_zip_contient_docx(self, app, job_with_docx_data):
        pytest.importorskip("docx")
        from transcria.config import get_config
        from transcria.exports.package_builder import PackageBuilder
        from transcria.jobs.models import Job, JobState

        with app.app_context():
            cfg = get_config()
            job = Job(id=job_with_docx_data, owner_id="admin", title="Test DOCX E2E",
                      state=JobState.CREATED.value)
            builder = PackageBuilder(cfg)
            result = builder.build_package(job)

        assert "zip_path" in result
        with zipfile.ZipFile(result["zip_path"], "r") as zf:
            names = zf.namelist()
            docx_files = [n for n in names if n.endswith(".docx")]
            assert docx_files, f"Aucun .docx dans le ZIP. Contenu: {names}"


class TestPackageLocalFreshness:
    """§5.3 : en backend fichiers LOCAL, le ZIP servi n'est jamais plus vieux que
    les artefacts sources — reconstruit au téléchargement s'il est périmé (même
    honnêteté que le backend pg)."""

    def test_zip_perime_est_reconstruit_au_download(self, admin_client, app, job_with_docx_data):
        import os
        import time

        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem

        with app.app_context():
            fs = JobFilesystem(get_config()["storage"]["jobs_dir"], job_with_docx_data)
        zip_path = fs.job_dir / "exports" / f"transcrIA_job_{job_with_docx_data}.zip"

        # 1er téléchargement : construit le zip (absent → périmé par définition).
        assert admin_client.get(f"/api/jobs/{job_with_docx_data}/download/package").status_code == 200
        assert zip_path.is_file()

        # Vieillir le zip puis modifier un artefact source (édition SRT simulée).
        old = time.time() - 3600
        os.utime(zip_path, (old, old))
        stale_mtime = zip_path.stat().st_mtime_ns
        (fs.job_dir / "metadata" / "transcription_corrigee.srt").write_text(
            "1\n00:00:00,000 --> 00:00:02,000\nTexte corrigé après coup.\n", encoding="utf-8")

        r = admin_client.get(f"/api/jobs/{job_with_docx_data}/download/package")

        assert r.status_code == 200
        assert zip_path.stat().st_mtime_ns > stale_mtime      # reconstruit, pas servi périmé
        import io
        import zipfile
        names = zipfile.ZipFile(io.BytesIO(r.data)).namelist()
        assert any(n.endswith(".srt") for n in names)

    def test_zip_a_jour_nest_pas_reconstruit(self, admin_client, app, job_with_docx_data):
        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem

        with app.app_context():
            fs = JobFilesystem(get_config()["storage"]["jobs_dir"], job_with_docx_data)
        zip_path = fs.job_dir / "exports" / f"transcrIA_job_{job_with_docx_data}.zip"

        assert admin_client.get(f"/api/jobs/{job_with_docx_data}/download/package").status_code == 200
        fresh_mtime = zip_path.stat().st_mtime_ns

        assert admin_client.get(f"/api/jobs/{job_with_docx_data}/download/package").status_code == 200
        assert zip_path.stat().st_mtime_ns == fresh_mtime     # zip frais servi tel quel

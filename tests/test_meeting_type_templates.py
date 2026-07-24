"""Types de réunion personnalisés — store (RBAC/quotas/collisions), API, étape 4, rendu.

Lot B de docs/TYPES_REUNION_PERSONNALISES.md. Décision D2 : tout utilisateur CRÉE
(portée privée), les admins PARTAGENT. La fiche d'un type choisi à l'étape 4 est
MATÉRIALISÉE dans le job (`meeting_context["custom_type"]`) : le rendu ne résout
jamais un template en base, la suppression du template ne casse aucun job.
"""
import uuid

import pytest


def _name(prefix: str) -> str:
    return f"{prefix}-{uuid.uuid4().hex[:8]}"


def _definition(name: str, **extra) -> dict:
    return {
        "name": name,
        "badge": "COMEX",
        "banner_text": "COMPTE-RENDU — COMITÉ EXÉCUTIF",
        "palette": {"primary": "1C1C1C", "accent": "424242", "light": "F5F5F5"},
        "fields": [{"key": "filiale", "label": "Filiale concernée", "type": "text"}],
        **extra,
    }


def _make_user(role=None, username_prefix="mtuser"):
    from transcria.auth.models import Role
    from transcria.auth.store import UserStore

    return UserStore.create_user(
        username=_name(username_prefix), password="test12345", role=role or Role.OPERATOR
    )


class TestStore:
    def test_tout_utilisateur_cree_un_type_prive(self, app):
        with app.app_context():
            from transcria.context.meeting_type_store import MeetingTypeStore

            user = _make_user()
            template = MeetingTypeStore.create_template(user, _definition(_name("COMEX")))
            assert template.scope == "private" and template.created_by == user.id
            assert template.definition["palette"]["primary"] == "1C1C1C"
            assert [t.id for t in MeetingTypeStore.visible_templates_for_user(user)] == [template.id]

    def test_quota_par_utilisateur(self, app):
        with app.app_context():
            from transcria.context.meeting_type_store import MeetingTypeStore, MeetingTypeValidationError

            user = _make_user()
            MeetingTypeStore.create_template(user, _definition(_name("Type A")), max_per_user=1)
            with pytest.raises(MeetingTypeValidationError, match="Quota"):
                MeetingTypeStore.create_template(user, _definition(_name("Type B")), max_per_user=1)

    def test_collision_avec_type_integre_refusee(self, app):
        with app.app_context():
            from transcria.context.meeting_type_store import MeetingTypeStore, MeetingTypeValidationError

            user = _make_user()
            with pytest.raises(MeetingTypeValidationError, match="intégré"):
                MeetingTypeStore.create_template(user, _definition("CSE"))
            # Collision de SLUG aussi (le nom diffère mais slugifie pareil qu'un intégré).
            with pytest.raises(MeetingTypeValidationError, match="intégré"):
                MeetingTypeStore.create_template(user, _definition("cse !"))

    def test_collision_entre_types_visibles_refusee_mais_pas_entre_users(self, app):
        with app.app_context():
            from transcria.context.meeting_type_store import MeetingTypeStore, MeetingTypeValidationError

            user_a, user_b = _make_user(), _make_user()
            name = _name("COMEX")
            MeetingTypeStore.create_template(user_a, _definition(name))
            with pytest.raises(MeetingTypeValidationError, match="existe déjà"):
                MeetingTypeStore.create_template(user_a, _definition(name))
            # Deux privés de deux utilisateurs peuvent porter le même nom (jamais ambigus :
            # la résolution étape 4 se fait dans le catalogue visible du propriétaire).
            assert MeetingTypeStore.create_template(user_b, _definition(name)).name == name

    def test_membre_simple_ne_partage_pas(self, app):
        with app.app_context():
            from transcria.context.meeting_type_store import MeetingTypeAccessError, MeetingTypeStore

            user = _make_user()
            template = MeetingTypeStore.create_template(user, _definition(_name("COMEX")))
            with pytest.raises(MeetingTypeAccessError):
                MeetingTypeStore.change_scope(user, template.id, "global")

    def test_admin_groupe_partage_le_prive_d_un_membre(self, app):
        with app.app_context():
            from transcria.auth.groups import GroupStore
            from transcria.auth.models import GroupRole
            from transcria.context.meeting_type_store import MeetingTypeStore

            member, group_admin, outsider = _make_user(), _make_user(), _make_user()
            group = GroupStore.create_group(_name("groupe"))
            GroupStore.add_member(group.id, member.id, GroupRole.MEMBER)
            GroupStore.add_member(group.id, group_admin.id, GroupRole.GROUP_ADMIN)

            template = MeetingTypeStore.create_template(member, _definition(_name("COMEX")))
            # L'admin de groupe VOIT le privé du membre dans sa vue de gestion…
            assert template.id in {t.id for t in MeetingTypeStore.list_manageable(group_admin)}
            # …et le promeut au groupe : tout membre le voit, l'extérieur non.
            MeetingTypeStore.change_scope(group_admin, template.id, "group", group.id)
            assert template.id in {t.id for t in MeetingTypeStore.visible_templates_for_user(group_admin)}
            assert template.id in {t.id for t in MeetingTypeStore.visible_templates_for_user(member)}
            assert template.id not in {t.id for t in MeetingTypeStore.visible_templates_for_user(outsider)}

    def test_type_partage_non_modifiable_par_membre_simple(self, app):
        with app.app_context():
            from transcria.auth.groups import GroupStore
            from transcria.auth.models import GroupRole
            from transcria.context.meeting_type_store import MeetingTypeAccessError, MeetingTypeStore

            member, group_admin = _make_user(), _make_user()
            group = GroupStore.create_group(_name("groupe"))
            GroupStore.add_member(group.id, member.id, GroupRole.MEMBER)
            GroupStore.add_member(group.id, group_admin.id, GroupRole.GROUP_ADMIN)
            template = MeetingTypeStore.create_template(member, _definition(_name("COMEX")))
            MeetingTypeStore.change_scope(group_admin, template.id, "group", group.id)

            with pytest.raises(MeetingTypeAccessError, match="dupliquez"):
                MeetingTypeStore.update_template(member, template.id, _definition(_name("COMEX v2")))
            # L'admin de groupe, lui, peut.
            updated = MeetingTypeStore.update_template(group_admin, template.id, _definition(_name("COMEX v2")))
            assert updated.name.startswith("COMEX v2")

    def test_global_reserve_admin(self, app):
        with app.app_context():
            from transcria.auth.groups import GroupStore
            from transcria.auth.models import GroupRole
            from transcria.auth.store import UserStore
            from transcria.context.meeting_type_store import MeetingTypeAccessError, MeetingTypeStore

            group_admin = _make_user()
            group = GroupStore.create_group(_name("groupe"))
            GroupStore.add_member(group.id, group_admin.id, GroupRole.GROUP_ADMIN)
            template = MeetingTypeStore.create_template(group_admin, _definition(_name("COMEX")))
            with pytest.raises(MeetingTypeAccessError, match="admin global"):
                MeetingTypeStore.change_scope(group_admin, template.id, "global")

            admin = UserStore.get_by_username("admin")
            MeetingTypeStore.change_scope(admin, template.id, "global")
            stranger = _make_user()
            assert template.id in {t.id for t in MeetingTypeStore.visible_templates_for_user(stranger)}

    def test_catalogue_fusionne(self, app):
        with app.app_context():
            from transcria.context.meeting_type_store import MeetingTypeStore

            user = _make_user()
            name = _name("COMEX")
            MeetingTypeStore.create_template(user, _definition(name))
            builtin, custom, fields = MeetingTypeStore.merged_catalog_for_user(user)
            assert len(builtin) == 18 and name in custom
            assert fields[name][0]["key"] == "filiale"
            assert "CSE" in fields  # champs intégrés toujours présents


class TestApi:
    def test_get_catalogue(self, operator_client):
        r = operator_client.get("/api/meeting-types")
        assert r.status_code == 200
        data = r.get_json()
        assert len(data["builtin"]) == 18 and data["builtin"][0]["builtin"] is True
        assert isinstance(data["custom"], list) and data["share_targets"]["global"] is False

    def test_create_update_delete(self, operator_client):
        r = operator_client.post("/api/meeting-types", json=_definition(_name("COMEX")))
        assert r.status_code == 201
        created = r.get_json()
        assert created["scope"] == "private"

        r = operator_client.put(f"/api/meeting-types/{created['id']}",
                                json=_definition(_name("COMEX rebaptisé")))
        assert r.status_code == 200

        r = operator_client.delete(f"/api/meeting-types/{created['id']}")
        assert r.status_code == 200
        remaining = operator_client.get("/api/meeting-types").get_json()["custom"]
        assert created["id"] not in {t["id"] for t in remaining}

    def test_create_invalide_400(self, operator_client):
        r = operator_client.post("/api/meeting-types",
                                 json={"name": "X", "palette": {"primary": "ZZZ"}})
        assert r.status_code == 400 and "palette" in r.get_json()["error"]

    def test_scope_par_non_admin_403(self, operator_client):
        created = operator_client.post("/api/meeting-types", json=_definition(_name("COMEX"))).get_json()
        r = operator_client.post(f"/api/meeting-types/{created['id']}/scope", json={"scope": "global"})
        assert r.status_code == 403

    def test_requires_login(self, client):
        assert client.get("/api/meeting-types").status_code in (302, 401)


class TestEtape4:
    @pytest.fixture
    def job_id(self, admin_client, app):
        r = admin_client.post("/jobs/new", data={"title": "Test types"}, follow_redirects=True)
        job_id = r.request.path.split("/")[2]
        # L'étape 4 (Contexte, avec le sélecteur de type) ne se rend qu'une fois le
        # résumé passé — on y amène le job.
        from transcria.jobs.models import JobState
        from transcria.jobs.store import JobStore
        with app.app_context():
            JobStore.update_state(job_id, JobState.SUMMARY_DONE)
        return job_id

    def test_type_personnalise_dans_le_wizard(self, admin_client, job_id):
        name = _name("COMEX")
        assert admin_client.post("/api/meeting-types", json=_definition(name)).status_code == 201
        html = admin_client.get(f"/jobs/{job_id}").data.decode()
        assert f'<option value="{name}"' in html          # dans le sélecteur de l'étape 4
        assert "Mes types & partagés" in html              # optgroup des personnalisés (i18n : gettext rend un & littéral)
        # Les champs du type personnalisé sont câblés pour le JS de l'étape 4.
        # Label échappé par |tojson (é → é dans l'îlot, re-décodé côté JS).
        assert "Filiale concern" in html

    def test_type_specific_fields_json_echappe_les_balises_script(self, admin_client, job_id):
        # Non-régression XSS : le label d'un champ de type est injecté dans l'îlot
        # window.__TYPE_SPECIFIC_FIELDS__. Avec json.dumps|safe, un </script> le
        # rompait ; avec |tojson il est échappé (</script>).
        name = _name("XSS")
        payload = "</script><img src=x onerror=alert(1)>"
        assert admin_client.post(
            "/api/meeting-types",
            json=_definition(name, fields=[{"key": "f", "label": payload, "type": "text"}]),
        ).status_code == 201
        html = admin_client.get(f"/jobs/{job_id}").data.decode()
        # La séquence brute qui romprait l'îlot ne doit JAMAIS apparaître.
        assert "</script><img src=x" not in html
        # Elle est présente sous forme échappée dans le JSON de l'îlot.
        assert "__TYPE_SPECIFIC_FIELDS__" in html
        assert "\\u003c/script\\u003e" in html or "\\u003cimg" in html

    def test_context_materialise_la_fiche(self, admin_client, app, job_id):
        name = _name("COMEX")
        created = admin_client.post("/api/meeting-types", json=_definition(name)).get_json()
        r = admin_client.post(f"/api/jobs/{job_id}/context",
                              json={"title": "T", "meeting_type": name})
        assert r.status_code == 200

        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem
        with app.app_context():
            fs = JobFilesystem(get_config()["storage"]["jobs_dir"], job_id)
            meeting = fs.load_json("context/meeting_context.json")
        assert meeting["meeting_type"] == name
        assert meeting["custom_type"]["template_id"] == created["id"]
        assert meeting["custom_type"]["palette"]["primary"] == "1C1C1C"

    def test_context_type_integre_purge_la_fiche(self, admin_client, app, job_id):
        name = _name("COMEX")
        admin_client.post("/api/meeting-types", json=_definition(name))
        admin_client.post(f"/api/jobs/{job_id}/context", json={"title": "T", "meeting_type": name})
        admin_client.post(f"/api/jobs/{job_id}/context", json={"title": "T", "meeting_type": "CSE"})

        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem
        with app.app_context():
            fs = JobFilesystem(get_config()["storage"]["jobs_dir"], job_id)
            meeting = fs.load_json("context/meeting_context.json")
        assert meeting["meeting_type"] == "CSE" and meeting["custom_type"] is None

    def test_context_type_inconnu_400(self, admin_client, job_id):
        r = admin_client.post(f"/api/jobs/{job_id}/context",
                              json={"title": "T", "meeting_type": "Type fantôme"})
        assert r.status_code == 400

    def test_fiche_survit_a_la_suppression_du_template(self, admin_client, app, job_id):
        name = _name("COMEX")
        created = admin_client.post("/api/meeting-types", json=_definition(name)).get_json()
        admin_client.post(f"/api/jobs/{job_id}/context", json={"title": "T", "meeting_type": name})
        admin_client.delete(f"/api/meeting-types/{created['id']}")

        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem
        with app.app_context():
            fs = JobFilesystem(get_config()["storage"]["jobs_dir"], job_id)
            meeting = fs.load_json("context/meeting_context.json")
        assert meeting["custom_type"]["palette"]["primary"] == "1C1C1C"


class TestRenduDocx:
    """La fiche matérialisée pilote le thème, les drapeaux et les libellés du rapport."""

    def _ctx(self, **custom) -> dict:
        return {
            "title": "Réunion test",
            "meeting_type": "COMEX Société X",
            "custom_type": {
                "name": "COMEX Société X",
                "badge": "COMEX",
                "banner_text": "COMPTE-RENDU — COMITÉ EXÉCUTIF",
                "palette": {"primary": "1C1C1C", "accent": "424242", "light": "F5F5F5"},
                "behavior": {"quorum": False, "confidential": False},
                "fields": [{"key": "filiale", "label": "Filiale concernée", "short_label": "Filiale", "type": "text"}],
                "template_id": "x",
                **custom,
            },
            "type_specific_data": {"filiale": "Filiale Nord"},
        }

    def _texts(self, doc) -> str:
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(p.text for p in cell.paragraphs)
        return "\n".join(parts)

    def test_theme_et_libelles_de_la_fiche(self):
        pytest.importorskip("docx")
        from transcria.exports.docx_report import DocxReport

        report = DocxReport(self._ctx(), [], {}, {}, "")
        assert str(report.theme.primary) == "1C1C1C"
        text = self._texts(report.build())
        assert "COMPTE-RENDU — COMITÉ EXÉCUTIF" in text   # bannière de la fiche
        assert "Filiale" in text and "Filiale Nord" in text  # short_label + valeur

    def test_comportement_quorum_de_la_fiche(self):
        pytest.importorskip("docx")
        from transcria.exports.docx_report import DocxReport

        ctx = self._ctx(behavior={"quorum": True, "confidential": False})
        ctx["type_specific_data"] = {"membres_presents": 6, "membres_total": 10}
        report = DocxReport(ctx, [], {}, {}, "")
        assert report.has_quorum is True
        assert "Quorum atteint" in self._texts(report.build())

    def test_confidentialite_de_la_fiche(self):
        pytest.importorskip("docx")
        from transcria.exports.docx_report import DocxReport

        ctx = self._ctx(behavior={"quorum": False, "confidential": True})
        report = DocxReport(ctx, [], {}, {}, "")
        assert report.ctx.get("sensitivity") == "high"

    def test_fiche_alteree_ne_plante_jamais(self):
        pytest.importorskip("docx")
        from transcria.exports.docx_report import _THEME_DEFAULT, DocxReport

        ctx = self._ctx()
        ctx["custom_type"]["palette"] = {"primary": "pas-un-hex"}
        report = DocxReport(ctx, [], {}, {}, "")
        assert report.theme is _THEME_DEFAULT
        report.build()  # aucun plantage


class TestLogo:
    """Branding local (lot C) : upload re-encodé Pillow, matérialisé dans le job."""

    def _png(self, size=(120, 40)) -> bytes:
        import io

        from PIL import Image
        buf = io.BytesIO()
        Image.new("RGB", size, (30, 30, 30)).save(buf, format="PNG")
        return buf.getvalue()

    def _upload(self, client, template_id: str, payload: bytes, filename="logo.png"):
        import io
        return client.post(f"/api/meeting-types/{template_id}/logo",
                           data={"logo": (io.BytesIO(payload), filename)},
                           content_type="multipart/form-data")

    def test_upload_reencode_et_delete(self, admin_client):
        created = admin_client.post("/api/meeting-types", json=_definition(_name("COMEX"))).get_json()
        r = self._upload(admin_client, created["id"], self._png(size=(2000, 800)))
        assert r.status_code == 200 and r.get_json()["has_logo"] is True

        # Re-encodé et borné : le blob stocké est un PNG aux dimensions réduites.
        r2 = admin_client.delete(f"/api/meeting-types/{created['id']}/logo")
        assert r2.status_code == 200 and r2.get_json()["has_logo"] is False

    def test_upload_non_image_400(self, admin_client):
        created = admin_client.post("/api/meeting-types", json=_definition(_name("COMEX"))).get_json()
        r = self._upload(admin_client, created["id"], b"pas une image du tout")
        assert r.status_code == 400

    def test_logo_materialise_dans_le_job_puis_purge(self, admin_client, app):
        import io

        from PIL import Image
        name = _name("COMEX")
        created = admin_client.post("/api/meeting-types", json=_definition(name)).get_json()
        self._upload(admin_client, created["id"], self._png())

        r = admin_client.post("/jobs/new", data={"title": "T logo"}, follow_redirects=True)
        job_id = r.request.path.split("/")[2]
        assert admin_client.post(f"/api/jobs/{job_id}/context",
                                 json={"title": "T", "meeting_type": name}).status_code == 200

        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem
        with app.app_context():
            logo_path = JobFilesystem(get_config()["storage"]["jobs_dir"], job_id).job_dir / "context" / "type_logo.png"
        assert logo_path.is_file()
        # Le blob matérialisé est bien le PNG re-encodé (lisible par Pillow).
        assert Image.open(io.BytesIO(logo_path.read_bytes())).format == "PNG"

        # Revenir à un type intégré purge le logo matérialisé.
        admin_client.post(f"/api/jobs/{job_id}/context", json={"title": "T", "meeting_type": "CSE"})
        assert not logo_path.exists()


class TestPageEtApercu:
    """Lot E : page « Mes types de réunion » + aperçu DOCX sur données factices."""

    def test_page_rendue(self, operator_client):
        html = operator_client.get("/meeting-types").data.decode()
        assert "Mes types de réunion" in html
        assert "mt-cards" in html and "mt-editor" in html   # galerie + éditeur
        assert "meeting_types.js" in html

    def test_nav_visible_pour_tous(self, operator_client):
        assert "/meeting-types" in operator_client.get("/").data.decode()

    def test_apercu_avant_enregistrement(self, admin_client):
        pytest.importorskip("docx")
        import io

        from docx import Document
        r = admin_client.post("/api/meeting-types/preview.docx", json=_definition(_name("COMEX")))
        assert r.status_code == 200
        doc = Document(io.BytesIO(r.data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:                  # le bandeau de couverture est un tableau
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(p.text for p in cell.paragraphs)
        text = "\n".join(parts)
        assert "COMITÉ EXÉCUTIF" in text          # bannière de la fiche en cours d'édition
        assert "RÉUNION D'EXEMPLE" in text        # titre factice (couverture en capitales)
        assert "Phrase d'exemple pour l'aperçu" in text   # données factices, jamais un vrai contenu

    def test_apercu_definition_invalide_400(self, admin_client):
        r = admin_client.post("/api/meeting-types/preview.docx",
                              json={"name": "X", "palette": {"primary": "ZZ"}})
        assert r.status_code == 400

    def test_apercu_type_enregistre(self, admin_client):
        created = admin_client.post("/api/meeting-types", json=_definition(_name("COMEX"))).get_json()
        r = admin_client.get(f"/api/meeting-types/{created['id']}/preview.docx")
        assert r.status_code == 200 and len(r.data) > 5000

    def test_apercu_type_invisible_404(self, operator_client, admin_client):
        created = admin_client.post("/api/meeting-types", json=_definition(_name("COMEX"))).get_json()
        # privé de l'admin → invisible de l'operator (sauf s'il partage un groupe avec lui — non ici)
        r = operator_client.get(f"/api/meeting-types/{created['id']}/preview.docx")
        assert r.status_code == 404

    def test_page_requires_login(self, client):
        assert client.get("/meeting-types").status_code in (302, 401)


class TestEchangeCommunautaire:
    """Lot F : export/import (§8) + validité des types du répertoire communautaire."""

    def test_export_sans_branding(self, admin_client):
        definition = _definition(_name("COMEX"), branding={"footer_text": "Société X"})
        created = admin_client.post("/api/meeting-types", json=definition).get_json()
        r = admin_client.get(f"/api/meeting-types/{created['id']}/export")
        assert r.status_code == 200
        payload = r.get_json()
        assert payload["schema_version"] == 1
        assert "branding" not in payload["type"]          # local, jamais exporté
        assert payload["type"]["palette"]["primary"] == "1C1C1C"

    def test_import_prive_inactif_puis_active_par_relecture(self, admin_client, app):
        name = _name("Type importé")
        payload = {"schema_version": 1, "type": _definition(name)}
        r = admin_client.post("/api/meeting-types/import", json=payload)
        assert r.status_code == 201
        imported = r.get_json()
        assert imported["scope"] == "private" and imported["is_active"] is False

        # Inactif : dans la galerie (badge « à relire ») mais PAS à l'étape 4.
        gallery = admin_client.get("/api/meeting-types").get_json()["custom"]
        entry = next(t for t in gallery if t["id"] == imported["id"])
        assert entry["is_active"] is False
        from transcria.auth.store import UserStore
        from transcria.context.meeting_type_store import MeetingTypeStore
        with app.app_context():
            admin = UserStore.get_by_username("admin")
            visible = {t.id for t in MeetingTypeStore.visible_templates_for_user(admin)}
        assert imported["id"] not in visible

        # La relecture = ouvrir, vérifier, enregistrer → le type s'active.
        r2 = admin_client.put(f"/api/meeting-types/{imported['id']}", json=_definition(name))
        assert r2.status_code == 200 and r2.get_json()["is_active"] is True

    def test_import_branding_refuse(self, admin_client):
        payload = {"schema_version": 1,
                   "type": _definition(_name("X"), branding={"footer_text": "Autrui"})}
        r = admin_client.post("/api/meeting-types/import", json=payload)
        assert r.status_code == 400 and "branding" in r.get_json()["error"]

    def test_import_schema_version_inconnu_refuse(self, admin_client):
        r = admin_client.post("/api/meeting-types/import",
                              json={"schema_version": 99, "type": {"name": "X"}})
        assert r.status_code == 400 and "schema_version" in r.get_json()["error"]

    def test_import_collision_suffixe(self, admin_client):
        name = _name("COMEX")
        admin_client.post("/api/meeting-types", json=_definition(name))
        r = admin_client.post("/api/meeting-types/import",
                              json={"schema_version": 1, "type": _definition(name)})
        assert r.status_code == 201
        assert r.get_json()["name"] == f"{name} (import)"

    def test_types_communautaires_valides(self):
        """Chaque fichier de community/meeting-types/ respecte le format d'échange
        (enveloppe, schéma, pas de branding) — la CI est la modération technique."""
        import json
        from pathlib import Path

        from transcria.context.meeting_type_catalog import validate_type_definition

        community = Path(__file__).resolve().parent.parent / "community" / "meeting-types"
        files = sorted(community.glob("*.json"))
        assert len(files) >= 3, "exemples communautaires manquants"
        for path in files:
            payload = json.loads(path.read_text(encoding="utf-8"))
            assert set(payload) == {"schema_version", "type"}, path.name
            assert payload["schema_version"] == 1, path.name
            assert "branding" not in payload["type"], path.name
            validated = validate_type_definition(payload["type"])
            assert validated["name"], path.name

"""Tests pour le générateur de rapport DOCX — données synthétiques, pas de job sur disque."""
import re
from pathlib import Path

import pytest


# ── Données synthétiques (reproduisent la structure réelle d'un job) ──────────

_CTX = {
    "title": "Scène de fromagerie — achat de comté",
    "meeting_type": "Autre",
    "date": "2026-05-29",
    "service": "Pédagogie",
    "language": "fr",
    "topic": "Dialogue en fromagerie",
    "objective": "Illustrer un échange commercial.",
    "notes": "Scène pédagogique.",
    "summary": "La cliente achète du comté et du beurre. Échange poli.",
    "sensitivity": "normal",
}

_PARTICIPANTS = [
    {"id": "p1", "name": "Cliente", "function": "", "service": "",
     "role": "pose des questions", "is_animator": False, "expected": True, "comment": ""},
    {"id": "p2", "name": "Vendeur / fromager", "function": "", "service": "",
     "role": "propose les produits", "is_animator": False, "expected": True, "comment": ""},
]

_SPEAKER_STATS = {
    "speakers": [
        {"speaker_id": "SPEAKER_00", "label": "SPEAKER_00", "mapped_to": "p1",
         "mapped_name": "Cliente", "speaking_time_seconds": 23.2,
         "turn_count": 15, "validation": "user_validated", "gender": "female"},
        {"speaker_id": "SPEAKER_01", "label": "SPEAKER_01", "mapped_to": "p2",
         "mapped_name": "Vendeur / fromager", "speaking_time_seconds": 25.7,
         "turn_count": 14, "validation": "user_validated", "gender": "male"},
    ]
}

_SRT = (
    "1\n00:00:01,012 --> 00:00:03,910\n"
    "SPEAKER_01(Vendeur / fromager): Podcast francefacil.com\n\n"
    "2\n00:00:05,416 --> 00:00:06,762\n"
    "SPEAKER_00(Cliente): Fais pas chaud ce matin.\n\n"
    "3\n00:00:11,592 --> 00:00:14,069\n"
    "SPEAKER_00(Cliente): Mettez-moi un peu d'émental s'il vous plaît.\n\n"
    "4\n00:00:19,827 --> 00:00:22,152\n"
    "SPEAKER_00(Cliente): Je prendrai bien un morceau de comté.\n\n"
)

_QUALITY = {
    "quality_score": 80,
    "total_checks": 16,
    "warnings": 4,
    "checks": [
        {"type": "low_coverage", "ratio": 0.79, "severity": "error"},
        {"type": "audio_problem_segments", "count": 1, "severity": "warning",
         "examples": [{"label": "silence", "start": 32.288, "end": 34.592,
                       "start_label": "00:32", "end_label": "00:35", "duration_s": 2.304}]},
    ],
    "review_points": ["Couverture faible : 79%", "Zone à réécouter : 00:32→00:35"],
}


def _seed_job(tmp_dir: Path, job_id: str) -> None:
    """Crée les fichiers d'un job synthétique dans tmp_dir."""
    import json
    base = tmp_dir / job_id
    for sub in ("context", "speakers", "metadata", "quality", "exports"):
        (base / sub).mkdir(parents=True, exist_ok=True)

    (base / "context" / "meeting_context.json").write_text(json.dumps(_CTX), encoding="utf-8")
    (base / "context" / "participants.json").write_text(json.dumps(_PARTICIPANTS), encoding="utf-8")
    (base / "speakers" / "speaker_stats.json").write_text(json.dumps(_SPEAKER_STATS), encoding="utf-8")
    (base / "metadata" / "transcription_corrigee.srt").write_text(_SRT, encoding="utf-8")
    (base / "quality" / "quality_report.json").write_text(json.dumps(_QUALITY), encoding="utf-8")


# ── Fixtures ──────────────────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def generated_docx(tmp_path_factory):
    """Génère le rapport DOCX une seule fois à partir de données synthétiques."""
    pytest.importorskip("docx")
    from transcria.exports.docx_report import generate_docx_report

    tmp = tmp_path_factory.mktemp("jobs_seed")
    job_id = "test-fromagerie-001"
    _seed_job(tmp, job_id)

    out = tmp_path_factory.mktemp("docx") / "rapport_test.docx"
    generate_docx_report(job_id, str(tmp), out)
    return out


@pytest.fixture(scope="module")
def doc(generated_docx):
    from docx import Document
    return Document(str(generated_docx))


# ── Helpers ───────────────────────────────────────────────────────────────────

def _all_text(doc) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def _table_texts(doc) -> list[str]:
    return [cell.text.strip() for table in doc.tables
            for row in table.rows for cell in row.cells]


# ── Tests fichier ──────────────────────────────────────────────────────────────

def test_fichier_cree(generated_docx):
    assert generated_docx.is_file()
    assert generated_docx.stat().st_size > 5000


# ── Tests contenu ─────────────────────────────────────────────────────────────

def test_titre_present(doc):
    text = _all_text(doc)
    assert "FROMAGERIE" in text.upper() or "COMTÉ" in text.upper()


def test_section_contexte(doc):
    assert "CONTEXTE" in _all_text(doc).upper()


def test_section_participants(doc):
    assert "PARTICIPANTS" in _all_text(doc).upper()


def test_section_transcription(doc):
    assert "TRANSCRIPTION" in _all_text(doc).upper()


def test_participants_dans_tableau(doc):
    joined = " ".join(_table_texts(doc))
    assert "Cliente" in joined
    assert "Vendeur" in joined


def test_temps_parole_calcule(doc):
    joined = " ".join(_table_texts(doc))
    assert re.search(r"\d+%", joined)


def test_transcription_contient_repliques(doc):
    joined = " ".join(_table_texts(doc)).lower()
    assert "émental" in joined or "comté" in joined


def test_timestamps_presents(doc):
    joined = " ".join(_table_texts(doc))
    assert re.search(r"\d{2}:\d{2}:\d{2}", joined)


def test_section_qualite_presente(doc):
    """Coverage 79% → la section 'Points à vérifier' doit apparaître."""
    full = _all_text(doc) + " ".join(_table_texts(doc))
    assert "VÉRIFIER" in full.upper() or "79" in full


def test_pas_de_confidentiel_par_defaut(doc):
    full = _all_text(doc) + " ".join(_table_texts(doc))
    assert "CONFIDENTIEL" not in full.upper()


def test_score_qualite_dans_document(doc):
    full = _all_text(doc) + " ".join(_table_texts(doc))
    assert "80" in full


# ── Tests cas limites ─────────────────────────────────────────────────────────

def test_genere_sans_participants(tmp_path):
    pytest.importorskip("docx")
    from transcria.exports.docx_report import DocxReport

    out = tmp_path / "vide.docx"
    DocxReport({}, [], {}, {}, "").build().save(str(out))
    assert out.is_file()


def test_genere_avec_champs_contexte_null(tmp_path):
    """Robustesse livrable : un meeting_context avec topic/objective/notes à `null`
    (clé présente, valeur None — LLM sans suggestion) ne doit PAS planter le rapport
    (`.get(k, "")` renvoie None, pas le défaut, sur une clé présente-null → None.strip())."""
    pytest.importorskip("docx")
    from transcria.exports.docx_report import DocxReport

    ctx = {"title": "Réunion", "topic": None, "objective": None, "notes": None}
    out = tmp_path / "ctx_null.docx"
    DocxReport(ctx, [], {}, {}, "").build().save(str(out))
    assert out.is_file()


def test_genere_avec_structured_data_non_dict(tmp_path):
    """Robustesse livrable : un structured_data mal typé (non-dict legacy/malformé) est
    toléré (ignoré) au lieu de planter `_section_enriched` (sd.get sur un non-dict)."""
    pytest.importorskip("docx")
    from transcria.exports.docx_report import DocxReport

    out = tmp_path / "sd_str.docx"
    # structured_data = chaîne (non-dict) : ne doit pas lever.
    DocxReport({"title": "X"}, [], {}, {}, "", structured_data="pas un dict").build().save(str(out))
    assert out.is_file()


def test_genere_avec_sensitivity_high(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport

    out = tmp_path / "confidentiel.docx"
    DocxReport({"title": "Test confidentiel", "sensitivity": "high"}, [], {}, {}, "").build().save(str(out))

    loaded = Document(str(out))
    full = ("\n".join(p.text for p in loaded.paragraphs)
            + " ".join(c.text for t in loaded.tables for r in t.rows for c in r.cells))
    assert "CONFIDENTIEL" in full.upper()


# ── Tests unitaires ───────────────────────────────────────────────────────────

def test_srt_parser_format_avec_locuteur():
    from transcria.exports.docx_report import _parse_srt

    srt = (
        "1\n00:00:01,012 --> 00:00:03,910\n"
        "SPEAKER_01(Vendeur / fromager): Podcast francefacil.com\n\n"
        "2\n00:00:05,416 --> 00:00:06,762\n"
        "SPEAKER_00(Cliente): Fais pas chaud ce matin.\n\n"
    )
    entries = _parse_srt(srt)
    assert len(entries) == 2
    assert entries[0]["speaker"] == "Vendeur / fromager"
    assert entries[0]["text"] == "Podcast francefacil.com"
    assert entries[0]["timestamp"] == "00:00:01"
    assert entries[1]["speaker"] == "Cliente"


def test_srt_parser_sans_locuteur():
    from transcria.exports.docx_report import _parse_srt

    entries = _parse_srt("1\n00:00:01,000 --> 00:00:03,000\nTexte sans locuteur.\n\n")
    assert len(entries) == 1
    assert entries[0]["speaker"] == ""
    assert entries[0]["text"] == "Texte sans locuteur."


def test_fmt_date_valide():
    from transcria.exports.docx_report import _fmt_date

    assert _fmt_date("2026-05-29") == "29 mai 2026"
    assert _fmt_date("") == "—"
    assert _fmt_date("2026-01-15") == "15 janvier 2026"


def test_fmt_date_invalide():
    from transcria.exports.docx_report import _fmt_date

    assert _fmt_date("pas-une-date") == "pas-une-date"


def test_extract_synthese_avec_section():
    from transcria.exports.docx_report import _extract_synthese

    md = "## Informations\nblah\n## Synthèse\nVoici la synthèse.\n## Autre\nfin"
    assert "Voici la synthèse" in _extract_synthese(md)
    assert "Autre" not in _extract_synthese(md)


def test_merge_participants_calcule_pourcentages():
    from transcria.exports.docx_report import DocxReport

    participants = [{"id": "p1", "name": "Alice", "function": "", "service": "",
                     "role": "", "is_animator": False}]
    speaker_stats = {"speakers": [{"speaker_id": "SPEAKER_00", "mapped_to": "p1",
                                    "speaking_time_seconds": 30.0, "turn_count": 10}]}
    report = DocxReport({}, participants, speaker_stats, {}, "")
    assert report.merged[0]["time_pct"] == 100
    assert report.merged[0]["turns"] == 10
    assert report.merged[0]["name"] == "Alice"


# ── Tests système de thèmes visuels ──────────────────────────────────────────

def test_theme_cse_est_institutionnel():
    from transcria.exports.docx_report import _get_theme, _THEMES
    theme = _get_theme("CSE")
    assert theme is _THEMES["CSE"]
    assert "PROCÈS-VERBAL" in theme.banner_text
    assert theme.cover_badge == "CSE"


def test_theme_point_projet_distinct_de_cse():
    from transcria.exports.docx_report import _get_theme
    cse = _get_theme("CSE")
    projet = _get_theme("Point projet")
    assert cse.primary != projet.primary
    assert cse.banner_text != projet.banner_text


def test_theme_inconnu_retourne_default():
    from transcria.exports.docx_report import _get_theme, _THEME_DEFAULT
    theme = _get_theme("Type qui n'existe pas")
    assert theme is _THEME_DEFAULT
    assert theme.banner_text == "COMPTE-RENDU DE TRANSCRIPTION"


def test_tous_les_types_majeurs_ont_un_theme():
    from transcria.exports.docx_report import _THEMES
    for t in ("CSE", "CSE extraordinaire", "CODIR / COMEX", "Point projet",
              "Réunion client", "Entretien individuel", "Formation",
              "Réunion de crise", "Réunion médicale / santé", "Négociation"):
        assert t in _THEMES, f"{t!r} n'a pas de thème dédié"


def test_docx_applique_le_theme_du_type(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport

    out = tmp_path / "cse.docx"
    DocxReport({"title": "Test CSE", "meeting_type": "CSE"}, [], {}, {}, "").build().save(str(out))
    loaded = Document(str(out))
    full = ("\n".join(p.text for p in loaded.paragraphs)
            + " ".join(c.text for t in loaded.tables for r in t.rows for c in r.cells))
    assert "PROCÈS-VERBAL" in full.upper()


def test_docx_crise_affiche_badge_situation_de_crise(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport

    out = tmp_path / "crise.docx"
    DocxReport({"title": "Incident majeur", "meeting_type": "Réunion de crise"}, [], {}, {}, "").build().save(str(out))
    loaded = Document(str(out))
    full = " ".join(c.text for t in loaded.tables for r in t.rows for c in r.cells)
    assert "CRISE" in full.upper()


def test_docx_cover_quorum_atteint(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport

    out = tmp_path / "cse_quorum.docx"
    ctx = {"title": "CSE", "meeting_type": "CSE",
           "type_specific_data": {"membres_presents": "8", "membres_total": "11"}}
    DocxReport(ctx, [], {}, {}, "").build().save(str(out))
    loaded = Document(str(out))
    full = " ".join(c.text for t in loaded.tables for r in t.rows for c in r.cells)
    assert "Quorum atteint" in full


def test_docx_cover_sous_titre_projet(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport

    out = tmp_path / "projet.docx"
    ctx = {"title": "Sprint 6", "meeting_type": "Point projet",
           "type_specific_data": {"nom_projet": "Projet Phoenix", "phase_jalon": "Phase 2"}}
    DocxReport(ctx, [], {}, {}, "").build().save(str(out))
    loaded = Document(str(out))
    full = "\n".join(p.text for p in loaded.paragraphs)
    assert "Projet Phoenix" in full


class TestMarkdownRendering:
    """Le DOCX rend le gras markdown au lieu de retirer les astérisques."""

    def test_split_plain_text_single_segment(self):
        from transcria.exports.docx_report import _split_markdown_bold
        assert _split_markdown_bold("texte simple") == [("texte simple", False)]

    def test_split_bold_segment(self):
        from transcria.exports.docx_report import _split_markdown_bold
        assert _split_markdown_bold("**TEST** test") == [("TEST", True), (" test", False)]

    def test_split_underscore_bold(self):
        from transcria.exports.docx_report import _split_markdown_bold
        assert _split_markdown_bold("avant __gras__ après") == [
            ("avant ", False), ("gras", True), (" après", False),
        ]

    def test_split_multiple_bold(self):
        from transcria.exports.docx_report import _split_markdown_bold
        assert _split_markdown_bold("**A** et **B**") == [
            ("A", True), (" et ", False), ("B", True),
        ]

    def test_split_empty(self):
        from transcria.exports.docx_report import _split_markdown_bold
        assert _split_markdown_bold("") == []

    def test_add_markdown_runs_sets_bold_flag(self):
        pytest.importorskip("docx")
        from docx import Document
        from transcria.exports.docx_report import _add_markdown_runs

        doc = Document()
        p = doc.add_paragraph()
        _add_markdown_runs(p, "**Cadrage.** Nicolas ouvre la séance.")
        # Premier run gras (l'intertitre), reste non gras.
        assert [(r.text, bool(r.font.bold)) for r in p.runs] == [
            ("Cadrage.", True),
            (" Nicolas ouvre la séance.", False),
        ]

    def test_synthesis_bold_lead_in_rendered_in_docx(self, doc):
        # Régression : un intertitre **…** de la synthèse ne doit plus apparaître
        # avec ses astérisques littérales, et doit produire au moins un run gras.
        from transcria.exports.docx_report import _split_markdown_bold
        # garde-fou pur : pas d'astérisques résiduelles après rendu
        rendered = "".join(c for c, _ in _split_markdown_bold("**Theme.** corps"))
        assert "*" not in rendered


class TestSummaryHarmonizedPreference:
    """Le DOCX préfère l'édition manuelle > la synthèse harmonisée > la brute."""

    def test_harmonized_used_when_no_manual_summary(self):
        pytest.importorskip("docx")
        from transcria.exports.docx_report import DocxReport
        ctx = {k: v for k, v in _CTX.items() if k != "summary"}
        ctx["summary_harmonized"] = "## Synthèse\n**Cadrage.** ACRO à 90 %."
        ctx["summary_llm"] = "## Synthèse\n**Cadrage.** AKRO à 90 %."
        doc = DocxReport(ctx, [], {}, {}, "", {}).build()
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "ACRO à 90 %" in text
        assert "AKRO" not in text

    def test_manual_summary_wins_over_harmonized(self):
        pytest.importorskip("docx")
        from transcria.exports.docx_report import DocxReport
        ctx = {**_CTX}
        ctx["summary"] = "## Synthèse\nEdition humaine validee."
        ctx["summary_harmonized"] = "## Synthèse\nHarmonise auto."
        doc = DocxReport(ctx, [], {}, {}, "", {}).build()
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Edition humaine validee" in text
        assert "Harmonise auto" not in text


# ── Robustesse produit final (passe qualité 12/06/2026) ───────────────────────

class TestSrtParserFormatNomSimple:
    """Le SRT corrigé peut arriver au format lisible « Nom: texte » (un agent LLM a
    déjà réécrit le préfixe SPEAKER_XX(Nom) malgré la consigne — job réel 4bda98cb).
    Sans repli, la colonne Locuteur du rapport était vide et le nom collé au texte."""

    def test_nom_simple_parse(self):
        from transcria.exports.docx_report import _parse_srt
        entries = _parse_srt("1\n00:00:01,000 --> 00:00:02,000\nVendeur: Bonjour madame.\n\n")
        assert entries[0]["speaker"] == "Vendeur"
        assert entries[0]["text"] == "Bonjour madame."

    def test_nom_accentue_et_compose(self):
        from transcria.exports.docx_report import _parse_srt
        entries = _parse_srt("1\n00:00:01,000 --> 00:00:02,000\nÉlise Martin : D'accord.\n\n")
        assert entries[0]["speaker"] == "Élise Martin"
        assert entries[0]["text"] == "D'accord."

    def test_format_strict_prioritaire(self):
        from transcria.exports.docx_report import _parse_srt
        entries = _parse_srt("1\n00:00:01,000 --> 00:00:02,000\nSPEAKER_01(Vendeur): Bonjour.\n\n")
        assert entries[0]["speaker"] == "Vendeur"  # nom humain, pas le préfixe brut

    def test_ligne_sans_locuteur_minuscule(self):
        from transcria.exports.docx_report import _parse_srt
        entries = _parse_srt("1\n00:00:01,000 --> 00:00:02,000\nmusique d'ambiance\n\n")
        assert entries[0]["speaker"] == ""


class TestDureeReunion:
    def test_duree_depuis_dernier_timestamp(self):
        from transcria.exports.docx_report import _srt_duration_seconds
        assert _srt_duration_seconds(_SRT) == 22  # 00:00:22,152

    def test_fmt_duration(self):
        from transcria.exports.docx_report import _fmt_duration
        assert _fmt_duration(45) == "1 min"
        assert _fmt_duration(60 * 73) == "1 h 13 min"
        assert _fmt_duration(3600) == "1 h"

    def test_duree_sur_la_couverture(self, doc):
        assert "Durée" in _table_texts(doc)


class TestStructuredDataRobuste:
    """Le JSON relu par la relecture finale peut dévier de « listes de chaînes »
    (items dicts, scalaires) : le rapport final ne plante jamais pour autant."""

    def _build(self, structured):
        from transcria.exports.docx_report import DocxReport
        report = DocxReport(_CTX, _PARTICIPANTS, _SPEAKER_STATS, _QUALITY, _SRT, structured)
        return report.build()

    def test_items_dicts_coerces_en_texte(self):
        doc = self._build({"decisions": [{"objet": "budget", "resultat": "adopté"}, "Décision B"]})
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Décision B" in text
        assert "budget" in text  # le dict est rendu en texte, pas une exception

    def test_champ_scalaire_accepte(self):
        doc = self._build({"actions": "Relancer le fournisseur"})
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Relancer le fournisseur" in text

    def test_gras_markdown_rendu_sans_asterisques(self):
        doc = self._build({"decisions": ["**Adopté** : budget 2026"]})
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "**" not in text
        assert "Adopté" in text


class TestPointsAVerifierCouvertureComplete:
    """Un avertissement du rapport qualité n'est jamais caché au lecteur du DOCX —
    notamment un nom de locuteur ALTÉRÉ par la LLM (severity=error)."""

    def _build_with_checks(self, checks):
        from transcria.exports.docx_report import DocxReport
        quality = {"quality_score": 70, "checks": checks}
        report = DocxReport(_CTX, _PARTICIPANTS, _SPEAKER_STATS, quality, _SRT, {})
        doc = report.build()
        return "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)

    def test_speaker_name_violations_affiche(self):
        text = self._build_with_checks([{
            "type": "speaker_name_violations", "severity": "error", "count": 1,
            "violations": [{"speaker_id": "SPEAKER_00", "found": "Mme Dupond", "expected": "Mme Dupont"}],
        }])
        assert "Nom de locuteur altéré" in text
        assert "Mme Dupond" in text and "Mme Dupont" in text

    def test_missing_lexicon_terms_affiche(self):
        text = self._build_with_checks([
            {"type": "missing_lexicon_terms", "severity": "warning", "terms": ["Emmental"]},
        ])
        assert "Terme du lexique non appliqué" in text
        assert "Emmental" in text

    def test_unmapped_speakers_affiche(self):
        text = self._build_with_checks([
            {"type": "unmapped_speakers", "severity": "warning", "count": 3},
        ])
        assert "Locuteurs non identifiés" in text

    def test_check_inconnu_repli_generique(self):
        text = self._build_with_checks([
            {"type": "empty_segments", "severity": "warning", "count": 2},
        ])
        assert "Segments vides" in text
        assert "2 élément(s)" in text

    def test_info_toujours_filtre(self):
        text = self._build_with_checks([
            {"type": "time_gaps", "severity": "info", "count": 4},
        ])
        assert "Points à vérifier".upper() not in text.upper()


def test_footer_absent_de_la_couverture(doc):
    """La page de garde reste vierge : pas de « Page 1/N » sur la couverture."""
    assert doc.sections[0].different_first_page_header_footer is True

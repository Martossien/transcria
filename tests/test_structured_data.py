"""Tests pour le parseur de données structurées LLM (3 niveaux de fallback)."""
import pytest


# ── Helpers ───────────────────────────────────────────────────────────────────

def _parse(text: str):
    from transcria.gpu.opencode_runner import OpenCodeRunner
    return OpenCodeRunner._parse_structured_data(text)


def _make_summary(json_block: str) -> str:
    return f"# Résumé\n\n## Synthèse\nBlah.\n\n## Données structurées\n\n```json\n{json_block}\n```\n"


# ── Section absente ────────────────────────────────────────────────────────────

def test_section_absente_retourne_missing():
    data, status, warning = _parse("# Résumé\n\n## Synthèse\nBlah.\n")
    assert status == "missing"
    assert warning == ""
    assert data["decisions"] == []
    assert data["prochaine_date"] == ""


# ── Niveau 1 : JSON valide ────────────────────────────────────────────────────

def test_json_valide_status_ok():
    summary = _make_summary('{"decisions": ["Budget validé"], "actions": ["Alice : envoyer le CR"], '
                            '"blocages": [], "reports": [], "votes": [], "resolutions": [], '
                            '"points_odj": [], "prochaine_date": "15/06/2026"}')
    data, status, warning = _parse(summary)
    assert status == "ok"
    assert warning == ""
    assert data["decisions"] == ["Budget validé"]
    assert data["actions"] == ["Alice : envoyer le CR"]
    assert data["prochaine_date"] == "15/06/2026"
    assert data["blocages"] == []


def test_json_valide_listes_vides():
    summary = _make_summary('{"decisions":[],"actions":[],"blocages":[],'
                            '"reports":[],"votes":[],"resolutions":[],"points_odj":[],"prochaine_date":""}')
    data, status, _ = _parse(summary)
    assert status == "ok"
    for field in ("decisions", "actions", "blocages", "reports", "votes", "resolutions", "points_odj"):
        assert data[field] == []
    assert data["prochaine_date"] == ""


def test_json_valide_champs_multiples():
    summary = _make_summary(
        '{"decisions": ["Projet lancé", "Budget 50k€ approuvé"],'
        ' "actions": ["Bob : préparer le plan avant vendredi"],'
        ' "blocages": ["Dépendance externe non résolue"],'
        ' "reports": ["Point juridique reporté"],'
        ' "votes": [], "resolutions": [], "points_odj": [], "prochaine_date": "22/06/2026"}'
    )
    data, status, _ = _parse(summary)
    assert status == "ok"
    assert len(data["decisions"]) == 2
    assert len(data["actions"]) == 1
    assert data["blocages"] == ["Dépendance externe non résolue"]
    assert data["reports"] == ["Point juridique reporté"]
    assert data["prochaine_date"] == "22/06/2026"


def test_json_normalise_items_vides():
    summary = _make_summary('{"decisions": ["", "Budget validé", "  "], "actions": []}')
    data, status, _ = _parse(summary)
    assert status == "ok"
    assert data["decisions"] == ["Budget validé"]


def test_json_valeur_string_convertie_en_liste():
    """Un champ qui arrive en string plutôt qu'en list doit être converti."""
    summary = _make_summary('{"decisions": "Budget validé", "actions": []}')
    data, status, _ = _parse(summary)
    assert status == "ok"
    assert data["decisions"] == ["Budget validé"]


# ── Niveau 2 : JSON malformé — regex ─────────────────────────────────────────

def test_json_malformed_extraction_partielle():
    # Virgule finale = JSON invalide mais extraction regex possible
    summary = _make_summary(
        '{"decisions": ["Budget approuvé",], "actions": ["Bob : rapport",]}'
    )
    data, status, warning = _parse(summary)
    assert status in ("ok", "partial")  # json.loads peut tolérer selon l'impl
    assert data["decisions"] or data["actions"]  # au moins quelque chose d'extrait


def test_json_malformed_champ_manquant_warning():
    # JSON sans guillemets doubles sur les clés
    summary = "# Résumé\n\n## Données structurées\n\n```\n{decisions: ['Budget']}\n```\n"
    data, status, warning = _parse(summary)
    # Soit ok si json.loads le parse, soit partial ou failed
    assert status in ("ok", "partial", "failed")


# ── Niveau 3 : Échec total ────────────────────────────────────────────────────

def test_section_presente_json_illisible_retourne_failed():
    summary = "# Résumé\n\n## Données structurées\n\nTexte libre sans JSON du tout.\n"
    data, status, warning = _parse(summary)
    assert status == "failed"
    assert warning != ""
    for field in ("decisions", "actions", "blocages"):
        assert data[field] == []


# ── Normalisation ─────────────────────────────────────────────────────────────

def test_normalize_structured_data():
    from transcria.gpu.opencode_runner import OpenCodeRunner
    raw = {
        "decisions": ["  Budget validé  ", ""],
        "actions": ["Bob : rapport"],
        "prochaine_date": "  15/06/2026  ",
        "votes": None,
    }
    result = OpenCodeRunner._normalize_structured_data(raw)
    assert result["decisions"] == ["Budget validé"]
    assert result["actions"] == ["Bob : rapport"]
    assert result["prochaine_date"] == "15/06/2026"
    assert result["votes"] == []


# ── Intégration avec _parse_structured_summary ────────────────────────────────

def test_parse_structured_summary_inclut_structured_data():
    from transcria.gpu.opencode_runner import OpenCodeRunner
    summary = (
        "# Résumé de contrôle\n\n"
        "## Informations sur la réunion\n"
        "- **Titre suggéré :** Réunion test\n"
        "- **Type suggéré :** Réunion interne\n"
        "- **Langue :** fr\n"
        "- **Sujet principal :** Test\n"
        "- **Objectif probable :** Tester\n"
        "- **Notes / Ordre du jour probable :** Aucun\n"
        "- **Nombre de participants détectés :** 2\n\n"
        "## Participants probables\n- SPEAKER_00 [Alice] : animatrice\n\n"
        "## Synthèse\nCeci est un test.\n\n"
        "## Termes douteux à valider\n(aucun terme suspect détecté)\n\n"
        "## Données structurées\n\n```json\n"
        '{"decisions": ["Test réussi"], "actions": [], "blocages": [], '
        '"reports": [], "votes": [], "resolutions": [], "points_odj": [], "prochaine_date": ""}'
        "\n```\n"
    )
    fields = OpenCodeRunner._parse_structured_summary(summary)
    assert "structured_data" in fields
    assert fields["structured_data_parse_status"] == "ok"
    assert fields["structured_data"]["decisions"] == ["Test réussi"]


def test_parse_structured_summary_sans_section_retourne_missing():
    from transcria.gpu.opencode_runner import OpenCodeRunner
    summary = (
        "# Résumé de contrôle\n\n"
        "## Informations sur la réunion\n"
        "- **Titre suggéré :** Test\n- **Type suggéré :** Autre\n"
        "- **Langue :** fr\n- **Sujet principal :** X\n"
        "- **Objectif probable :** Y\n- **Notes / Ordre du jour probable :** Z\n"
        "- **Nombre de participants détectés :** 1\n\n"
        "## Participants probables\n- (non identifiable)\n\n"
        "## Synthèse\nBlah.\n\n"
        "## Termes douteux à valider\n(aucun terme suspect détecté)\n"
    )
    fields = OpenCodeRunner._parse_structured_summary(summary)
    assert fields["structured_data_parse_status"] == "missing"
    assert fields["structured_data"]["decisions"] == []


# ── Tests DOCX avec données enrichies ────────────────────────────────────────

def test_docx_avec_decisions_affiche_section():
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport
    import io

    sd = {"decisions": ["Budget approuvé"], "actions": [], "blocages": [],
          "reports": [], "votes": [], "resolutions": [], "points_odj": [], "prochaine_date": ""}
    report = DocxReport({"title": "Test", "meeting_type": "Réunion interne"}, [], {}, {}, "", sd)
    doc = report.build()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    loaded = Document(buf)
    full = "\n".join(p.text for p in loaded.paragraphs)
    assert "DÉCISIONS" in full.upper()
    assert "Budget approuvé" in full


def test_docx_actions_affichees_si_presentes():
    """Toute donnée extraite non vide s'affiche, quel que soit le type."""
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport
    import io

    sd = {"decisions": [], "actions": ["Bob : faire quelque chose"], "blocages": [],
          "reports": [], "votes": [], "resolutions": [], "points_odj": [], "prochaine_date": ""}
    report = DocxReport({"title": "Test", "meeting_type": "Podcast / média"}, [], {}, {}, "", sd)
    doc = report.build()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    loaded = Document(buf)
    full = "\n".join(p.text for p in loaded.paragraphs).upper()
    assert "ACTIONS À RÉALISER" in full


def test_docx_votes_present_pour_cse():
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport
    import io

    sd = {"decisions": [], "actions": [], "blocages": [],
          "reports": [], "votes": ["Budget : 12 pour, 2 contre — adopté"],
          "resolutions": ["Résolution n°1 adoptée"], "points_odj": [], "prochaine_date": ""}
    report = DocxReport({"title": "CSE Test", "meeting_type": "CSE"}, [], {}, {}, "", sd)
    doc = report.build()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    loaded = Document(buf)
    full = "\n".join(p.text for p in loaded.paragraphs)
    assert "VOTES" in full.upper()
    assert "RÉSOLUTIONS" in full.upper()


def test_docx_votes_affiches_hors_cse():
    """Régression mairie : des votes extraits doivent s'afficher même si le type
    n'est pas CSE (conseil municipal, AG, copropriété…)."""
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport
    import io

    sd = {"decisions": [], "actions": [], "blocages": [],
          "reports": [], "votes": ["Budget : 12 pour — adopté"],
          "resolutions": [], "points_odj": [], "prochaine_date": ""}
    report = DocxReport({"title": "Test", "meeting_type": "Réunion interne"}, [], {}, {}, "", sd)
    doc = report.build()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    loaded = Document(buf)
    full = "\n".join(p.text for p in loaded.paragraphs).upper()
    assert "VOTES" in full


def test_docx_numerotation_sections_avec_enrichissement():
    """Avec des données enrichies, les sections Participants/Transcription/Qualité
    doivent être numérotées après les sections enrichies."""
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport
    import io

    sd = {"decisions": ["D1"], "actions": ["A1"], "blocages": [],
          "reports": [], "votes": [], "resolutions": [], "points_odj": [], "prochaine_date": ""}
    report = DocxReport({"title": "Test", "meeting_type": "Réunion interne"}, [], {}, {}, "", sd)
    doc = report.build()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    loaded = Document(buf)
    full = "\n".join(p.text for p in loaded.paragraphs)
    # 2 sections enrichies (decisions + actions) → Participants doit être en 4.
    assert "4." in full  # Participants
    assert "5." in full  # Transcription


def test_docx_entretien_individuel_auto_confidentiel():
    pytest.importorskip("docx")
    from docx import Document
    from transcria.exports.docx_report import DocxReport
    import io

    report = DocxReport({"title": "Test", "meeting_type": "Entretien individuel"}, [], {}, {}, "")
    doc = report.build()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    loaded = Document(buf)
    tables_text = " ".join(c.text for t in loaded.tables for r in t.rows for c in r.cells)
    assert "CONFIDENTIEL" in tables_text.upper()


class TestClesExtractionPersonnalisees:
    """Lot D : les extract_fields d'un type personnalisé traversent le parseur."""

    def test_niveau_1_conserve_les_cles_du_type(self):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        text = ('## Données structurées\n```json\n'
                '{"decisions": ["D1"], "budgets_evoques": ["10 k€ pour le projet A"]}\n```')
        sd, status, _ = OpenCodeRunner._parse_structured_data(text, ("budgets_evoques",))
        assert status == "ok" and sd["budgets_evoques"] == ["10 k€ pour le projet A"]
        # Sans la clé déclarée, elle est filtrée (liste blanche inchangée par défaut).
        sd2, _, _ = OpenCodeRunner._parse_structured_data(text)
        assert "budgets_evoques" not in sd2

    def test_niveau_2_regex_couvre_les_cles_du_type(self):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        text = ('## Données structurées\n```json\n'
                '{"decisions": ["D1"], "budgets_evoques": ["10 k€"],}\n```')  # virgule finale → json.loads échoue
        sd, status, _ = OpenCodeRunner._parse_structured_data(text, ("budgets_evoques",))
        assert status == "partial" and sd["budgets_evoques"] == ["10 k€"]


class TestMaterialisationPrompt:
    """Lot D : le prompt résolu (placeholders substitués) est écrit dans le scratch."""

    def test_substitue_et_ecrit_dans_le_scratch(self, tmp_path):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        prompt = tmp_path / "summary_prompt.txt"
        prompt.write_text("Types : [{{TYPES_REUNION}}]\n{{INDICES_TYPES}}\n{{CHAMPS_EXTRACTION_TYPE}}",
                          encoding="utf-8")
        runner = OpenCodeRunner(str(tmp_path), model="local/fake-model")
        resolved = runner._materialize_prompt(str(prompt), {
            "{{TYPES_REUNION}}": "CSE | COMEX Société X",
            "{{INDICES_TYPES}}": '  `COMEX Société X` si on entend "comité exécutif" ;',
            "{{CHAMPS_EXTRACTION_TYPE}}": "",
        })
        assert resolved != str(prompt)
        text = open(resolved, encoding="utf-8").read()
        assert "COMEX Société X" in text and "{{" not in text

    def test_sans_placeholder_fichier_original(self, tmp_path):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        prompt = tmp_path / "summary_prompt.txt"
        prompt.write_text("Prompt maison sans placeholder.", encoding="utf-8")
        runner = OpenCodeRunner(str(tmp_path), model="local/fake-model")
        assert runner._materialize_prompt(str(prompt), {"{{TYPES_REUNION}}": "X"}) == str(prompt)

"""GOLDEN du rendu DOCX — filet posé AVANT l'extraction des helpers de docx_report.py (vague 0).

L'empreinte figée (tests/golden_docx_report.json) est le rendu COMPLET du job synthétique du
module de tests voisin : chaque paragraphe (style + texte) et chaque cellule de tableau, dans
l'ordre. Les tests de contenu existants vérifient des PRÉSENCES (« CONTEXTE apparaît ») ; ce
golden verrouille l'INTÉGRALITÉ — un helper de style déplacé qui perdrait un thème, une
bordure de section ou l'ordre d'un tableau rougirait ici et nulle part ailleurs.

Régénération après une évolution VOULUE du rendu :
    venv/bin/python -m pytest tests/test_docx_report_golden.py --force-regen  (cf. fixture)
    (ou supprimer tests/golden_docx_report.json puis relancer — il se réécrit et le diff git
    montre exactement ce qui a changé dans le document)
"""
from __future__ import annotations

import importlib.util
import json
import re
from pathlib import Path

import pytest

pytest.importorskip("docx")

GOLDEN_PATH = Path(__file__).parent / "golden_docx_report.json"


def _load_seed():
    """Réutilise le _seed_job du module de tests voisin — même document canonique, une seule
    source de données synthétiques."""
    path = Path(__file__).parent / "test_docx_report.py"
    spec = importlib.util.spec_from_file_location("_docx_seed_module", path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module._seed_job


_STAMPED_DATE = re.compile(r"\b\d{2}/\d{2}/\d{4}\b")


def _norm(text: str) -> str:
    """La couverture estampille la date de GÉNÉRATION (datetime.today) — normalisée, sinon le
    golden casserait chaque jour. Les dates métier (« 29 mai 2026 ») restent vérifiées."""
    return _STAMPED_DATE.sub("<DATE>", text)


def _fingerprint(doc) -> dict:
    return {
        "paragraphs": [[p.style.name, _norm(p.text)] for p in doc.paragraphs],
        "tables": [
            [[_norm(cell.text) for cell in row.cells] for row in table.rows]
            for table in doc.tables
        ],
        "sections": len(doc.sections),
    }


@pytest.fixture(scope="module")
def canonical_fingerprint(tmp_path_factory):
    from docx import Document

    from transcria.exports.docx_report import generate_docx_report

    seed = _load_seed()
    tmp = tmp_path_factory.mktemp("jobs_seed_golden")
    job_id = "test-fromagerie-001"
    seed(tmp, job_id)
    out = tmp_path_factory.mktemp("docx_golden") / "rapport_golden.docx"
    generate_docx_report(job_id, str(tmp), out)
    return _fingerprint(Document(str(out)))


def test_full_render_matches_the_golden(canonical_fingerprint):
    if not GOLDEN_PATH.exists():
        GOLDEN_PATH.write_text(
            json.dumps(canonical_fingerprint, ensure_ascii=False, indent=1), encoding="utf-8"
        )
        pytest.skip("golden généré (premier run) — relancer pour vérifier")
    golden = json.loads(GOLDEN_PATH.read_text(encoding="utf-8"))
    assert canonical_fingerprint["sections"] == golden["sections"]
    assert canonical_fingerprint["tables"] == golden["tables"]
    assert canonical_fingerprint["paragraphs"] == golden["paragraphs"]

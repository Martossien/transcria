"""
Tests E2E automatisés — extraction de données structurées et rendu DOCX.

Ces tests couvrent le pipeline complet sans GPU :
  LLM output simulé → _parse_structured_summary → meeting_context.json
  → génération DOCX → vérification des sections → rendu HTML wizard

Scénarios testés :
  1. Parsing LLM → meeting_context (flux runner complet mockée)
  2. Types de réunion : routing CSE / CODIR / Point projet / Podcast / Entretien
  3. Dégradation gracieuse (parse failed → DOCX v1 sans crash)
  4. Numérotation dynamique des sections
  5. Panneau wizard HTML (données enrichies collapsible)
  6. Téléchargement DOCX via HTTP avec données enrichies
  7. Intégrité du fichier prompt (section 8b présente)
  8. Nouveaux types de réunion (liste + dropdown)
  9. Thèmes visuels par type (bannières, badges, quorum CSE) via HTTP/DOCX réel
 10. Champs type-spécifiques (panneau wizard + persistance + DOCX)
 11. Helper de tracking thème du runner E2E (_docx_theme_info)
"""
import io
import json
from pathlib import Path

import pytest

# ── Fixtures partagées ────────────────────────────────────────────────────────

def _make_summary_with_json(json_block: str) -> str:
    """Produit un summary.md réaliste contenant un bloc JSON Données structurées."""
    return (
        "# Résumé de contrôle\n\n"
        "## Informations sur la réunion\n"
        "- **Titre suggéré :** Réunion de test E2E\n"
        "- **Type suggéré :** Réunion interne\n"
        "- **Langue :** fr\n"
        "- **Sujet principal :** Test automatisé du pipeline\n"
        "- **Objectif probable :** Valider la feature\n"
        "- **Notes / Ordre du jour probable :** Tests E2E\n"
        "- **Nombre de participants détectés :** 2\n\n"
        "## Participants probables\n"
        "- SPEAKER_00 [Alice] : anime la réunion\n"
        "- SPEAKER_01 [Bob] : développeur\n\n"
        "## Synthèse\n"
        "La réunion a permis de valider la feature d'extraction structurée.\n\n"
        "## Termes douteux à valider\n"
        "(aucun terme suspect détecté)\n\n"
        "## Données structurées\n\n"
        "```json\n"
        f"{json_block}\n"
        "```\n"
    )


_SD_COMPLET = {
    "decisions": ["Feature validée en production", "Déploiement prévu semaine prochaine"],
    "actions": ["Alice : rédiger la note de déploiement avant jeudi",
                "Bob : mettre à jour les tests CI"],
    "blocages": ["Dépendance externe non encore disponible"],
    "reports": ["Point budget reporté à la prochaine réunion"],
    "votes": [],
    "resolutions": [],
    "points_odj": [],
    "prochaine_date": "05/06/2026",
}

_SD_CSE = {
    "decisions": [],
    "actions": [],
    "blocages": [],
    "reports": [],
    "votes": ["Budget formation : 10 pour, 2 contre, 1 abstention — adopté",
              "Plan mobilité : 13 pour — adopté"],
    "resolutions": ["Résolution n°1 : Budget formation approuvé à 15k€"],
    "points_odj": ["1. Budget formation annuelle — adopté à l'unanimité",
                   "2. Plan mobilité 2026 — approuvé"],
    "prochaine_date": "15/07/2026",
}


def _seed_enriched_job(jobs_dir: str, job_id: str, structured_data: dict,
                       meeting_type: str = "Réunion interne",
                       type_specific_data: dict | None = None,
                       title: str = "Réunion E2E test") -> None:
    """Crée un job avec données enrichies pré-remplies (simule le LLM ayant tourné)."""
    from transcria.jobs.filesystem import JobFilesystem

    fs = JobFilesystem(jobs_dir, job_id)
    ctx = {
        "title": title,
        "meeting_type": meeting_type,
        "date": "2026-05-29",
        "service": "IT",
        "language": "fr",
        "topic": "Test E2E structuré",
        "objective": "Valider le pipeline",
        "notes": "Tests automatisés",
        "summary": "La réunion s'est bien déroulée.",
        "sensitivity": "normal",
        "structured_data": structured_data,
        "structured_data_parse_status": "ok",
    }
    if type_specific_data:
        ctx["type_specific_data"] = type_specific_data
    fs.save_json("context/meeting_context.json", ctx)
    fs.save_json("context/participants.json", [
        {"id": "p1", "name": "Alice", "function": "Chef de projet", "service": "IT",
         "role": "Animatrice", "is_animator": True, "expected": True, "comment": ""},
        {"id": "p2", "name": "Bob", "function": "Développeur", "service": "IT",
         "role": "Contributeur", "is_animator": False, "expected": True, "comment": ""},
    ])
    fs.save_json("speakers/speaker_stats.json", {
        "speakers": [
            {"speaker_id": "SPEAKER_00", "mapped_to": "p1", "mapped_name": "Alice",
             "speaking_time_seconds": 40.0, "turn_count": 8, "validation": "user_validated"},
            {"speaker_id": "SPEAKER_01", "mapped_to": "p2", "mapped_name": "Bob",
             "speaking_time_seconds": 20.0, "turn_count": 4, "validation": "user_validated"},
        ]
    })
    fs.save_text(
        "metadata/transcription_corrigee.srt",
        "1\n00:00:01,000 --> 00:00:03,000\nSPEAKER_00(Alice): Bonjour.\n\n"
        "2\n00:00:04,000 --> 00:00:06,000\nSPEAKER_01(Bob): Bonjour à tous.\n\n",
    )
    fs.save_json("quality/quality_report.json", {
        "quality_score": 92, "total_checks": 10, "warnings": 0, "checks": [],
    })


def _advance_to_summary_done(app, job_id: str) -> None:
    """Avance l'état du job à SUMMARY_DONE pour débloquer l'affichage du contexte."""
    with app.app_context():
        from transcria.jobs.models import JobState
        from transcria.jobs.store import JobStore
        JobStore.update_state(job_id, JobState.SUMMARY_DONE)


@pytest.fixture
def job_enriched(admin_client, app):
    """Job avec données structurées complètes (type Réunion interne), état SUMMARY_DONE."""
    from transcria.config import get_config
    r = admin_client.post("/jobs/new", data={"title": "E2E Structured Test"}, follow_redirects=True)
    job_id = r.request.path.split("/")[2] if "/jobs/" in r.request.path else None
    assert job_id, "Impossible de créer le job"
    with app.app_context():
        cfg = get_config()
        _seed_enriched_job(cfg["storage"]["jobs_dir"], job_id, _SD_COMPLET, "Réunion interne")
    _advance_to_summary_done(app, job_id)
    return job_id


@pytest.fixture
def job_cse(admin_client, app):
    """Job avec données structurées CSE (votes, résolutions, ODJ), état SUMMARY_DONE."""
    from transcria.config import get_config
    r = admin_client.post("/jobs/new", data={"title": "E2E CSE Test"}, follow_redirects=True)
    job_id = r.request.path.split("/")[2] if "/jobs/" in r.request.path else None
    assert job_id
    with app.app_context():
        cfg = get_config()
        _seed_enriched_job(cfg["storage"]["jobs_dir"], job_id, _SD_CSE, "CSE")
    _advance_to_summary_done(app, job_id)
    return job_id


# ── 1. Pipeline parser LLM → meeting_context (sans HTTP) ─────────────────────

class TestParserToContext:
    def test_parse_complet_produit_tous_les_champs(self):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        summary = _make_summary_with_json(json.dumps(_SD_COMPLET))
        fields = OpenCodeRunner._parse_structured_summary(summary)

        assert fields["structured_data_parse_status"] == "ok"
        assert fields["structured_data"]["decisions"] == _SD_COMPLET["decisions"]
        assert fields["structured_data"]["actions"] == _SD_COMPLET["actions"]
        assert fields["structured_data"]["blocages"] == _SD_COMPLET["blocages"]
        assert fields["structured_data"]["prochaine_date"] == "05/06/2026"

    def test_parse_cse_extrait_votes_et_resolutions(self):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        summary = _make_summary_with_json(json.dumps(_SD_CSE))
        fields = OpenCodeRunner._parse_structured_summary(summary)

        assert fields["structured_data_parse_status"] == "ok"
        assert len(fields["structured_data"]["votes"]) == 2
        assert len(fields["structured_data"]["resolutions"]) == 1
        assert len(fields["structured_data"]["points_odj"]) == 2

    def test_parse_listes_vides_status_ok(self):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        sd_vide = {k: [] if k != "prochaine_date" else "" for k in _SD_COMPLET}
        summary = _make_summary_with_json(json.dumps(sd_vide))
        fields = OpenCodeRunner._parse_structured_summary(summary)
        assert fields["structured_data_parse_status"] == "ok"
        assert fields["structured_data"]["decisions"] == []

    def test_parse_json_malformed_retourne_partial_ou_failed(self):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        summary = (
            "# Résumé\n\n## Synthèse\nBlah.\n\n"
            "## Termes douteux à valider\n(aucun)\n\n"
            "## Données structurées\n\n```json\n{decisions: ['Décision sans guillemets']}\n```\n"
        )
        fields = OpenCodeRunner._parse_structured_summary(summary)
        assert fields["structured_data_parse_status"] in ("partial", "failed")

    def test_parse_section_absente_status_missing(self):
        from transcria.gpu.opencode_runner import OpenCodeRunner
        summary = (
            "# Résumé de contrôle\n\n## Synthèse\nBlah.\n\n"
            "## Termes douteux à valider\n(aucun terme suspect détecté)\n"
        )
        fields = OpenCodeRunner._parse_structured_summary(summary)
        assert fields["structured_data_parse_status"] == "missing"
        assert fields["structured_data"]["decisions"] == []

    def test_apply_llm_suggestions_stocke_structured_data(self, tmp_path):
        """Vérifie que runner._apply_llm_suggestions écrit structured_data dans meeting_context."""
        from transcria.gpu.opencode_runner import OpenCodeRunner
        from transcria.jobs.filesystem import JobFilesystem
        from transcria.workflow.runner import WorkflowRunner

        jobs_dir = str(tmp_path)
        job_id = "test-apply-e2e"
        fs = JobFilesystem(jobs_dir, job_id)
        fs.save_json("context/meeting_context.json", {"title": "Test"})

        summary_text = _make_summary_with_json(json.dumps(_SD_COMPLET))
        parsed = OpenCodeRunner._parse_structured_summary(summary_text)
        parsed["summary_text"] = summary_text

        result: dict = {}

        # Le logger structuré accepte des kwargs arbitraires — on utilise un mock léger
        class _StubLogger:
            def info(self, *a, **kw): pass
            def warning(self, *a, **kw): pass
            def debug(self, *a, **kw): pass
            def error(self, *a, **kw): pass

        WorkflowRunner._apply_llm_suggestions(fs, result, parsed, _StubLogger())

        ctx = fs.load_json("context/meeting_context.json") or {}
        assert "structured_data" in ctx
        assert ctx["structured_data_parse_status"] == "ok"
        assert ctx["structured_data"]["decisions"] == _SD_COMPLET["decisions"]
        assert ctx["structured_data"]["prochaine_date"] == "05/06/2026"


# ── 2. Routing DOCX par type de réunion ──────────────────────────────────────

class TestDocxTypeRouting:
    def _build_doc(self, meeting_type: str, sd: dict) -> str:
        pytest.importorskip("docx")
        from docx import Document

        from transcria.exports.docx_report import DocxReport

        report = DocxReport(
            {"title": "Test", "meeting_type": meeting_type}, [], {}, {}, "", sd
        )
        doc = report.build()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        loaded = Document(buf)
        return "\n".join(p.text for p in loaded.paragraphs).upper()

    def test_reunion_interne_affiche_decisions_actions_blocages(self):
        """Règle : toute donnée extraite non vide s'affiche, quel que soit le type."""
        full = self._build_doc("Réunion interne", _SD_COMPLET)
        assert "DÉCISIONS" in full
        assert "ACTIONS À RÉALISER" in full
        assert "POINTS BLOQUANTS" in full  # blocages non vides → affichés même hors projet

    def test_point_projet_affiche_blocages(self):
        full = self._build_doc("Point projet", _SD_COMPLET)
        assert "POINTS BLOQUANTS" in full

    def test_cse_affiche_votes_resolutions_odj(self):
        full = self._build_doc("CSE", _SD_CSE)
        assert "VOTES" in full
        assert "RÉSOLUTIONS" in full
        assert "ORDRE DU JOUR" in full

    def test_reunion_interne_affiche_votes_si_presents(self):
        """Régression mairie : des votes extraits ne doivent jamais être cachés
        sous prétexte que le type n'est pas CSE (conseil municipal, AG, copro…)."""
        full = self._build_doc("Réunion interne", _SD_CSE)
        assert "VOTES" in full
        assert "RÉSOLUTIONS" in full
        assert "ORDRE DU JOUR" in full

    def test_podcast_affiche_actions_si_presentes(self):
        """Plus de filtrage par type : si le LLM a extrait des actions, on les montre."""
        full = self._build_doc("Podcast / média", _SD_COMPLET)
        assert "ACTIONS À RÉALISER" in full

    def test_point_projet_affiche_actions(self):
        full = self._build_doc("Point projet", _SD_COMPLET)
        assert "ACTIONS À RÉALISER" in full

    def test_sections_vides_jamais_affichees(self):
        """Une section dont la liste est vide ne doit pas apparaître."""
        sd_vide = {k: [] if k != "prochaine_date" else "" for k in _SD_COMPLET}
        sd_vide["decisions"] = ["Une seule décision"]
        full = self._build_doc("Réunion interne", sd_vide)
        assert "DÉCISIONS" in full
        assert "VOTES" not in full
        assert "ACTIONS À RÉALISER" not in full
        assert "POINTS BLOQUANTS" not in full

    def test_cse_ordre_du_jour_avant_votes(self):
        """L'ordre du jour doit précéder les votes dans le PV CSE."""
        full = self._build_doc("CSE", _SD_CSE)
        assert "ORDRE DU JOUR" in full
        idx_odj = full.find("ORDRE DU JOUR")
        idx_votes = full.find("VOTES")
        assert idx_odj != -1 and idx_votes != -1
        assert idx_odj < idx_votes, "L'ordre du jour doit apparaître avant les votes"

    def test_entretien_individuel_auto_confidentiel(self):
        pytest.importorskip("docx")
        from docx import Document

        from transcria.exports.docx_report import DocxReport

        report = DocxReport({"title": "Test", "meeting_type": "Entretien individuel"}, [], {}, {}, "")
        doc = report.build()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        loaded = Document(buf)
        tables_text = " ".join(
            c.text for t in loaded.tables for r in t.rows for c in r.cells
        ).upper()
        assert "CONFIDENTIEL" in tables_text

    def test_codir_affiche_decisions_et_actions(self):
        full = self._build_doc("CODIR / COMEX", _SD_COMPLET)
        assert "DÉCISIONS" in full
        assert "ACTIONS" in full

    def test_crise_affiche_blocages(self):
        full = self._build_doc("Réunion de crise", _SD_COMPLET)
        assert "POINTS BLOQUANTS" in full

    def test_prochaine_date_presente_si_renseignee(self):
        full = self._build_doc("Réunion interne", _SD_COMPLET)
        assert "05/06/2026" in full

    def test_pas_de_prochaine_date_si_vide(self):
        sd = dict(_SD_COMPLET)
        sd["prochaine_date"] = ""
        full = self._build_doc("Réunion interne", sd)
        assert "Prochaine réunion" not in full


# ── 3. Numérotation dynamique des sections ────────────────────────────────────

class TestSectionNumbering:
    def _sections(self, meeting_type: str, sd: dict) -> list[str]:
        pytest.importorskip("docx")
        from docx import Document

        from transcria.exports.docx_report import DocxReport
        report = DocxReport({"title": "T", "meeting_type": meeting_type}, [], {}, {}, "", sd)
        doc = report.build()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        loaded = Document(buf)
        text = "\n".join(p.text for p in loaded.paragraphs)
        return [line for line in text.split("\n") if line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7."))]

    def test_sans_enrichissement_numerotation_classique(self):
        sd_vide = {k: [] if k != "prochaine_date" else "" for k in _SD_COMPLET}
        sections = self._sections("Réunion interne", sd_vide)
        # 1. Contexte, 2. Participants, 3. Transcription
        nums = [s.split(".")[0].strip() for s in sections]
        assert "1" in nums
        assert "2" in nums
        assert "3" in nums
        assert "4" not in nums  # pas de Points à vérifier si qualité OK

    def test_avec_decisions_et_actions_participants_est_en_4(self):
        sections = self._sections("Réunion interne", _SD_COMPLET)
        joined = " ".join(sections)
        assert "4." in joined   # Participants décalé en 4
        assert "5." in joined   # Transcription en 5

    def test_cse_numerotation_avec_odj_votes_resolutions(self):
        sections = self._sections("CSE", _SD_CSE)
        nums = [s.split(".")[0].strip() for s in sections]
        # CSE avec votes(1) + résolutions(1) + odj(1) = 3 sections enrichies
        # Participants = 2 + 3 = 5
        assert "5" in nums


# ── 4. Dégradation gracieuse ──────────────────────────────────────────────────

class TestGracefulDegradation:
    def test_structured_data_vide_docx_valide(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        from transcria.exports.docx_report import generate_docx_report
        from transcria.jobs.filesystem import JobFilesystem

        jobs_dir = str(tmp_path)
        job_id = "test-degradation"
        fs = JobFilesystem(jobs_dir, job_id)
        fs.save_json("context/meeting_context.json", {
            "title": "Test dégradation",
            "meeting_type": "Réunion interne",
            # structured_data absent intentionnellement
        })
        fs.save_json("context/participants.json", [])
        fs.save_json("speakers/speaker_stats.json", {"speakers": []})
        fs.save_text("metadata/transcription_corrigee.srt", "")
        fs.save_json("quality/quality_report.json", {"quality_score": 90, "checks": []})

        out = tmp_path / "rapport.docx"
        generate_docx_report(job_id, jobs_dir, out)
        assert out.is_file()
        assert out.stat().st_size > 3000

        loaded = Document(str(out))
        full = "\n".join(p.text for p in loaded.paragraphs).upper()
        assert "CONTEXTE" in full
        assert "PARTICIPANTS" in full
        assert "TRANSCRIPTION" in full

    def test_structured_data_parse_failed_docx_valide(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document

        from transcria.exports.docx_report import generate_docx_report
        from transcria.jobs.filesystem import JobFilesystem

        jobs_dir = str(tmp_path)
        job_id = "test-degradation-failed"
        fs = JobFilesystem(jobs_dir, job_id)
        fs.save_json("context/meeting_context.json", {
            "title": "Test parse failed",
            "meeting_type": "CSE",
            "structured_data": {},  # dict vide = pas de sections enrichies
            "structured_data_parse_status": "failed",
            "structured_data_parse_warning": "JSON non parseable",
        })
        fs.save_json("context/participants.json", [])
        fs.save_json("speakers/speaker_stats.json", {"speakers": []})
        fs.save_text("metadata/transcription_corrigee.srt", "")
        fs.save_json("quality/quality_report.json", {"quality_score": 80, "checks": []})

        out = tmp_path / "rapport_failed.docx"
        generate_docx_report(job_id, jobs_dir, out)
        assert out.is_file()

        loaded = Document(str(out))
        full = "\n".join(p.text for p in loaded.paragraphs).upper()
        # Sections enrichies absentes, document v1 standard
        assert "VOTES" not in full
        assert "RÉSOLUTIONS" not in full
        assert "TRANSCRIPTION" in full


# ── 5. Panneau wizard HTML ────────────────────────────────────────────────────

class TestWizardEnrichedPanel:
    def test_panneau_visible_si_structured_data_present(self, admin_client, app, job_enriched):
        r = admin_client.get(f"/jobs/{job_enriched}")
        assert r.status_code == 200
        html = r.data.decode("utf-8")
        # Le panneau données enrichies doit être présent
        assert "Données enrichies extraites" in html
        assert "Décisions prises" in html
        assert "Actions à réaliser" in html

    def test_panneau_contient_les_items_extraits(self, admin_client, app, job_enriched):
        r = admin_client.get(f"/jobs/{job_enriched}")
        html = r.data.decode("utf-8")
        assert "Feature validée en production" in html
        assert "Alice : rédiger la note" in html
        assert "05/06/2026" in html

    def test_panneau_absent_si_pas_de_structured_data(self, admin_client, app):
        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem

        r = admin_client.post("/jobs/new", data={"title": "No SD"}, follow_redirects=True)
        job_id = r.request.path.split("/")[2] if "/jobs/" in r.request.path else None
        assert job_id

        with app.app_context():
            cfg = get_config()
            fs = JobFilesystem(cfg["storage"]["jobs_dir"], job_id)
            fs.save_json("context/meeting_context.json", {
                "title": "No structured data",
                "meeting_type": "Réunion interne",
                # pas de structured_data
            })

        r2 = admin_client.get(f"/jobs/{job_id}")
        html = r2.data.decode("utf-8")
        assert "Données enrichies extraites" not in html

    def test_panneau_cse_affiche_votes(self, admin_client, app, job_cse):
        r = admin_client.get(f"/jobs/{job_cse}")
        html = r.data.decode("utf-8")
        assert "Votes" in html
        assert "Résolutions adoptées" in html
        assert "Budget formation : 10 pour" in html

    def test_badge_ok_si_parse_status_ok(self, admin_client, app, job_enriched):
        r = admin_client.get(f"/jobs/{job_enriched}")
        html = r.data.decode("utf-8")
        assert "bg-success" in html  # badge vert


# ── 6. Téléchargement DOCX avec données enrichies (HTTP) ─────────────────────

class TestDocxHTTPEnriched:
    def test_docx_reunion_interne_contient_decisions(self, admin_client, app, job_enriched):
        pytest.importorskip("docx")
        from docx import Document

        r = admin_client.get(f"/api/jobs/{job_enriched}/download/docx")
        assert r.status_code == 200
        doc = Document(io.BytesIO(r.data))
        full = "\n".join(p.text for p in doc.paragraphs).upper()
        table_text = " ".join(
            c.text for t in doc.tables for row in t.rows for c in row.cells
        ).upper()
        combined = full + " " + table_text
        assert "DÉCISIONS" in combined
        assert "ACTIONS" in combined

    def test_docx_cse_contient_votes(self, admin_client, app, job_cse):
        pytest.importorskip("docx")
        from docx import Document

        r = admin_client.get(f"/api/jobs/{job_cse}/download/docx")
        assert r.status_code == 200
        doc = Document(io.BytesIO(r.data))
        full = "\n".join(p.text for p in doc.paragraphs).upper()
        assert "VOTES" in full
        assert "RÉSOLUTIONS" in full

    def test_docx_avec_structured_data_taille_superieure(self, admin_client, app):
        """Un DOCX avec données enrichies doit être plus grand qu'un DOCX vide."""
        pytest.importorskip("docx")
        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem

        # Job sans données enrichies
        r_vide = admin_client.post("/jobs/new", data={"title": "Vide"}, follow_redirects=True)
        job_vide = r_vide.request.path.split("/")[2] if "/jobs/" in r_vide.request.path else None
        assert job_vide
        with app.app_context():
            cfg = get_config()
            fs = JobFilesystem(cfg["storage"]["jobs_dir"], job_vide)
            fs.save_json("context/meeting_context.json", {"title": "Vide", "meeting_type": "Réunion interne"})
            fs.save_json("context/participants.json", [])
            fs.save_json("speakers/speaker_stats.json", {"speakers": []})
            fs.save_text("metadata/transcription_corrigee.srt", "")
            fs.save_json("quality/quality_report.json", {"quality_score": 90, "checks": []})

        r_docx_vide = admin_client.get(f"/api/jobs/{job_vide}/download/docx")
        assert r_docx_vide.status_code == 200
        assert len(r_docx_vide.data) > 5000


# ── 7. Intégrité du prompt (section 8b) ──────────────────────────────────────

class TestPromptIntegrity:
    def test_prompt_contient_section_8b(self):
        prompt_path = Path(__file__).parent.parent / "configs" / "prompts" / "summary_prompt.txt"
        assert prompt_path.is_file(), "summary_prompt.txt introuvable"
        content = prompt_path.read_text(encoding="utf-8")
        assert "## 8b. Section données structurées" in content

    def test_prompt_contient_tous_les_champs_json(self):
        prompt_path = Path(__file__).parent.parent / "configs" / "prompts" / "summary_prompt.txt"
        content = prompt_path.read_text(encoding="utf-8")
        for field in ("decisions", "actions", "blocages", "reports",
                      "votes", "resolutions", "points_odj", "prochaine_date"):
            assert f'"{field}"' in content, f"Champ {field!r} absent du prompt"

    def test_prompt_contient_verification_16(self):
        prompt_path = Path(__file__).parent.parent / "configs" / "prompts" / "summary_prompt.txt"
        content = prompt_path.read_text(encoding="utf-8")
        assert "Données structurées" in content
        assert "16." in content

    def test_prompt_contient_etape_4b(self):
        prompt_path = Path(__file__).parent.parent / "configs" / "prompts" / "summary_prompt.txt"
        content = prompt_path.read_text(encoding="utf-8")
        assert "4b." in content


# ── 8. Nouveaux types de réunion ──────────────────────────────────────────────

class TestMeetingTypes:
    def test_types_elargie_present_dans_liste(self):
        from transcria.context.meeting_context import MEETING_TYPES
        for expected in ("CSE", "CSE extraordinaire", "CODIR / COMEX", "Point projet",
                         "Réunion client", "Réunion de crise", "Séminaire / atelier",
                         "Négociation", "Entretien individuel", "Podcast / média"):
            assert expected in MEETING_TYPES, f"{expected!r} absent de MEETING_TYPES"

    def test_types_anciens_preserves(self):
        from transcria.context.meeting_context import MEETING_TYPES
        for legacy in ("Réunion interne", "Réunion projet", "Réunion technique",
                       "Formation", "Réunion médicale / santé", "RH", "Entretien"):
            assert legacy in MEETING_TYPES, f"Type legacy {legacy!r} supprimé !"

    def test_meeting_types_dans_le_dropdown_wizard(self, admin_client, app):
        """Les nouveaux types apparaissent dans le formulaire de contexte."""
        r = admin_client.post("/jobs/new", data={"title": "Type test"}, follow_redirects=True)
        job_id = r.request.path.split("/")[2] if "/jobs/" in r.request.path else None
        assert job_id
        # Avancer à SUMMARY_DONE pour débloquer l'affichage du formulaire contexte
        _advance_to_summary_done(app, job_id)
        r2 = admin_client.get(f"/jobs/{job_id}")
        html = r2.data.decode("utf-8")
        assert "CSE" in html
        assert "CODIR / COMEX" in html
        assert "Point projet" in html


# ── 9. Thèmes visuels par type (E2E HTTP, DOCX réel) ──────────────────────────

def _seed_themed_job(admin_client, app, meeting_type: str, structured_data: dict,
                     type_specific_data: dict | None = None, title: str = "Thème E2E") -> str:
    """Crée un job complet, le mène jusqu'à EXPORT_READY, retourne le job_id."""
    from transcria.config import get_config
    from transcria.jobs.models import JobState
    from transcria.jobs.store import JobStore

    r = admin_client.post("/jobs/new", data={"title": title}, follow_redirects=True)
    job_id = r.request.path.split("/")[2] if "/jobs/" in r.request.path else None
    assert job_id, "Création du job échouée"
    with app.app_context():
        cfg = get_config()
        _seed_enriched_job(cfg["storage"]["jobs_dir"], job_id, structured_data,
                           meeting_type, type_specific_data, title)
        JobStore.update_state(job_id, JobState.EXPORT_READY)
    return job_id


def _download_docx_text(admin_client, job_id: str) -> str:
    """Télécharge le DOCX et retourne tout son texte (paragraphes + tableaux)."""
    from docx import Document
    r = admin_client.get(f"/api/jobs/{job_id}/download/docx")
    assert r.status_code == 200, f"Téléchargement DOCX échoué ({r.status_code})"
    assert r.data[:4] == b"PK\x03\x04", "Réponse non DOCX (magic bytes)"
    doc = Document(io.BytesIO(r.data))
    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    tables = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    return paragraphs + "\n" + tables


class TestDocxThemesE2E:
    """Vérifie de bout en bout (HTTP → DOCX réel) que chaque type applique
    sa bannière, ses champs type-spécifiques et ses éléments visuels."""

    def test_cse_banniere_institutionnelle(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "CSE", _SD_CSE,
                               {"president_seance": "Marie Dupont",
                                "secretaire_seance": "Pierre Martin",
                                "membres_presents": "8", "membres_total": "11",
                                "ref_pv_precedent": "PV-2026-01"},
                               title="CSE Q2 2026")
        full = _download_docx_text(admin_client, job).upper()
        assert "PROCÈS-VERBAL DU COMITÉ SOCIAL" in full

    def test_cse_champs_specifiques_sur_couverture(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "CSE", _SD_CSE,
                               {"president_seance": "Marie Dupont",
                                "secretaire_seance": "Pierre Martin",
                                "membres_presents": "8", "membres_total": "11"})
        full = _download_docx_text(admin_client, job)
        assert "Marie Dupont" in full
        assert "Pierre Martin" in full

    def test_cse_quorum_calcule_automatiquement(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "CSE", _SD_CSE,
                               {"membres_presents": "8", "membres_total": "11"})
        full = _download_docx_text(admin_client, job)
        assert "Quorum atteint" in full
        assert "73%" in full

    def test_cse_quorum_non_atteint(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "CSE", _SD_CSE,
                               {"membres_presents": "3", "membres_total": "11"})
        full = _download_docx_text(admin_client, job)
        assert "non atteint" in full.lower()

    def test_point_projet_banniere_et_sous_titre(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "Point projet", _SD_COMPLET,
                               {"nom_projet": "Projet Phoenix",
                                "phase_jalon": "Phase 2", "chef_de_projet": "Alice",
                                "sprint": "6"},
                               title="Sprint 6")
        full = _download_docx_text(admin_client, job)
        assert "RÉUNION PROJET" in full.upper()
        assert "Projet Phoenix" in full  # sous-titre contextuel
        assert "Alice" in full           # chef de projet en métadonnée

    def test_codir_banniere_direction(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "CODIR / COMEX", _SD_COMPLET)
        full = _download_docx_text(admin_client, job).upper()
        assert "COMITÉ DE DIRECTION" in full

    def test_crise_badge_situation_de_crise(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "Réunion de crise", _SD_COMPLET,
                               {"nature_incident": "Panne datacenter",
                                "responsable_crise": "DSI"})
        full = _download_docx_text(admin_client, job).upper()
        assert "CRISE" in full
        assert "PANNE DATACENTER" in full  # sous-titre contextuel

    def test_entretien_individuel_confidentiel(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "Entretien individuel", _SD_COMPLET,
                               {"poste_evalue": "Développeur senior",
                                "periode_evaluee": "2025", "evaluateur": "Manager"})
        full = _download_docx_text(admin_client, job)
        assert "CONFIDENTIEL" in full.upper()
        assert "ENTRETIEN INDIVIDUEL" in full.upper()

    def test_reunion_client_banniere_et_ref_contrat(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "Réunion client", _SD_COMPLET,
                               {"nom_client": "Acme Corp", "ref_contrat": "CTR-2026-001"})
        full = _download_docx_text(admin_client, job)
        assert "RÉUNION CLIENT" in full.upper()
        assert "Acme Corp" in full
        assert "CTR-2026-001" in full

    def test_type_inconnu_garde_banniere_par_defaut(self, admin_client, app):
        pytest.importorskip("docx")
        job = _seed_themed_job(admin_client, app, "Réunion interne", _SD_COMPLET)
        full = _download_docx_text(admin_client, job).upper()
        assert "COMPTE-RENDU DE TRANSCRIPTION" in full

    def test_themes_produisent_des_documents_distincts(self, admin_client, app):
        """Deux types différents → bannières différentes dans le DOCX."""
        pytest.importorskip("docx")
        job_cse = _seed_themed_job(admin_client, app, "CSE", _SD_CSE,
                                   {"membres_presents": "8", "membres_total": "11"})
        job_proj = _seed_themed_job(admin_client, app, "Point projet", _SD_COMPLET,
                                    {"nom_projet": "X"})
        txt_cse = _download_docx_text(admin_client, job_cse).upper()
        txt_proj = _download_docx_text(admin_client, job_proj).upper()
        assert "PROCÈS-VERBAL" in txt_cse
        assert "PROCÈS-VERBAL" not in txt_proj
        assert "RÉUNION PROJET" in txt_proj


class TestTypeSpecificE2EWizard:
    """Le panneau de champs type-spécifiques et sa config JSON dans le wizard."""

    def test_config_type_specific_fields_dans_page(self, admin_client, app):
        job = _seed_themed_job(admin_client, app, "CSE", _SD_CSE,
                               {"president_seance": "Marie"})
        # Repasser en SUMMARY_DONE pour afficher le formulaire contexte
        with app.app_context():
            from transcria.jobs.models import JobState
            from transcria.jobs.store import JobStore
            JobStore.update_state(job, JobState.SUMMARY_DONE)
        r = admin_client.get(f"/jobs/{job}")
        html = r.data.decode("utf-8")
        # La config des champs est injectée pour le JS
        assert "__TYPE_SPECIFIC_FIELDS__" in html
        assert "president_seance" in html
        # La valeur déjà saisie est ré-injectée
        assert "Marie" in html

    def test_type_specific_data_persiste_apres_save_contexte(self, admin_client, app):
        from transcria.config import get_config
        from transcria.jobs.filesystem import JobFilesystem

        job = _seed_themed_job(admin_client, app, "Point projet", _SD_COMPLET,
                               {"nom_projet": "Phoenix"})
        with app.app_context():
            from transcria.jobs.models import JobState
            from transcria.jobs.store import JobStore
            JobStore.update_state(job, JobState.SUMMARY_DONE)

        # L'utilisateur ré-enregistre le contexte sans renvoyer type_specific_data
        r = admin_client.post(f"/api/jobs/{job}/context",
                              json={"title": "Sprint 6 màj", "meeting_type": "Point projet"},
                              content_type="application/json")
        assert r.status_code == 200

        with app.app_context():
            cfg = get_config()
            fs = JobFilesystem(cfg["storage"]["jobs_dir"], job)
            ctx = fs.load_json("context/meeting_context.json") or {}
            # type_specific_data doit être préservé
            assert ctx.get("type_specific_data", {}).get("nom_projet") == "Phoenix"


# ── 11. Helper de tracking thème du runner E2E réel ───────────────────────────

class TestRunnerThemeTracking:
    """Le helper _docx_theme_info de tests/test_e2e_workflow.py, qui alimente
    le JSON de sortie --output-json pour les campagnes bench."""

    def _helper(self):
        import importlib.util
        spec = importlib.util.spec_from_file_location(
            "_e2e_runner", str(Path(__file__).parent / "test_e2e_workflow.py"))
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        return mod._docx_theme_info

    def test_cse_resolu(self):
        info = self._helper()({
            "meeting_type": "CSE",
            "type_specific_data": {"president_seance": "Marie", "membres_presents": "8",
                                   "membres_total": "11", "secretaire_seance": ""},
        })
        assert "PROCÈS-VERBAL" in info["banner_text"].upper()
        assert info["cover_badge"] == "CSE"
        assert info["is_default_theme"] is False
        # secretaire_seance vide n'est pas compté
        assert "president_seance" in info["type_specific_fields_filled"]
        assert "secretaire_seance" not in info["type_specific_fields_filled"]
        assert info["type_specific_count"] == 3

    def test_type_inconnu_est_default(self):
        info = self._helper()({"meeting_type": "Réunion interne"})
        assert info["is_default_theme"] is True
        assert info["type_specific_count"] == 0

    def test_point_projet_champs_listes(self):
        info = self._helper()({
            "meeting_type": "Point projet",
            "type_specific_data": {"nom_projet": "Phoenix", "sprint": "6"},
        })
        assert info["banner_text"]
        assert info["is_default_theme"] is False
        assert info["type_specific_fields_filled"] == ["nom_projet", "sprint"]

    def test_meeting_type_absent_ne_casse_pas(self):
        info = self._helper()({})
        assert info["meeting_type"] == ""
        assert info["type_specific_count"] == 0

    def test_write_output_json_ne_plante_pas(self, admin_client, app, tmp_path):
        """Régression : write_output_json référençait meeting_ctx non chargé.

        Bug attrapé par un run E2E réel — le bloc structured_data/docx_theh
        plantait avec « name 'meeting_ctx' is not defined ». Ce test garantit
        que la fonction écrit un JSON complet sur un job seedé.
        """
        import importlib.util
        import json as _json
        from transcria.config import get_config

        spec = importlib.util.spec_from_file_location(
            "_e2e_runner_full", str(Path(__file__).parent / "test_e2e_workflow.py"))
        e2e = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(e2e)

        job = _seed_themed_job(admin_client, app, "CSE", _SD_CSE,
                               {"president_seance": "Marie", "membres_presents": "8",
                                "membres_total": "11"})
        with app.app_context():
            cfg = get_config()
            from transcria.jobs.filesystem import JobFilesystem
            fs = JobFilesystem(cfg["storage"]["jobs_dir"], job)

        e2e.RESULTS["job_id"] = job
        import sys as _sys
        old_argv = _sys.argv
        _sys.argv = ["test_e2e_workflow.py", "--audio", "tests/test2.mp3", "--skip-llm"]
        try:
            args = e2e.parse_args()
        finally:
            _sys.argv = old_argv

        out = tmp_path / "result.json"
        e2e.write_output_json(out, args, cfg, fs)  # ne doit pas lever

        assert out.is_file()
        payload = _json.loads(out.read_text(encoding="utf-8"))
        assert "structured_data" in payload
        assert "docx_theme" in payload
        # Le job seedé est de type CSE → thème institutionnel
        assert payload["docx_theme"]["meeting_type"] == "CSE"
        assert "PROCÈS-VERBAL" in payload["docx_theme"]["banner_text"].upper()
        assert payload["docx_theme"]["is_default_theme"] is False

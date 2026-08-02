"""Génère un rapport DOCX professionnel à partir des artefacts d'un job terminé."""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentT  # type pour annotations (docx.Document est une fabrique)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor

from transcria.context.meeting_type_catalog import (
    ORDERABLE_SECTIONS,
    confidential_types,
    localized_field_labels,
    localized_type_display,
    quorum_types,
)
from transcria.exports.docx_srt import (  # noqa: F401 — façade
    _parse_srt,
    _srt_duration_seconds,
)
from transcria.exports.docx_style import (  # noqa: F401 — façade
    _BLUE_DARK,
    _BLUE_LIGHT,
    _BLUE_MID,
    _DOCX_LABELS,
    _GREEN,
    _GREY_DARK,
    _GREY_LIGHT,
    _LANG_LABELS,
    _ORANGE,
    _RED,
    _THEME_DEFAULT,
    _THEMES,
    _WHITE,
    _YELLOW_BG,
    _add_markdown_runs,
    _add_num_pages_field,
    _add_page_number_field,
    _cell_bg,
    _cell_margins,
    _docx_labels,
    _DocxTheme,
    _fmt_date,
    _fmt_duration,
    _fmt_time,
    _get_theme,
    _para_bottom_border,
    _split_markdown_bold,
    _table_full_width,
    _table_no_borders,
    _table_thin_borders,
    _theme_from_definition,
)
from transcria.jobs.filesystem import JobFilesystem
from transcria.llm_tools.opencode_runner import summary_markers

# ── Palette ──────────────────────────────────────────────────────────────────
_CSE_TYPES: frozenset[str] = quorum_types()
_AUTO_CONFIDENTIEL: frozenset[str] = confidential_types()

# ── Parsing SRT ───────────────────────────────────────────────────────────────

_RENDER_SECTIONS = ("participants", "transcript", "quality")

# Unités ordonnables du registre de sections (lot C). « couverture » est fixe ;
# « contexte » et « pv » sont déplaçables mais JAMAIS désactivables (règle : une
# donnée extraite n'est jamais cachée). « synthese »/« champs_type » ne deviennent
# des sections autonomes que si un ordre les cite — sinon ils restent DANS le
# contexte (rendu historique, non-régression au pixel).
_ORDERABLE_SECTIONS = ORDERABLE_SECTIONS
_SECTION_ORDER_DEFAULT = ("contexte", "pv", "participants", "transcript", "quality")


def _sanitize_render_options(raw: object) -> dict:
    """Valide ``context/render_options.json`` — tout invalide est ignoré (le rendu ne casse jamais).

    Options v1 : ``theme`` (clé de ``_THEMES``, prime sur le meeting_type),
    ``sections`` (booléens ``participants``/``transcript``/``quality``) et
    ``order`` (liste d'unités de ``_ORDERABLE_SECTIONS``, clés inconnues ignorées).
    """
    if not isinstance(raw, dict):
        return {}
    out: dict = {}
    theme = raw.get("theme")
    if isinstance(theme, str) and theme in _THEMES:
        out["theme"] = theme
    sections = raw.get("sections")
    if isinstance(sections, dict):
        cleaned = {k: bool(v) for k, v in sections.items() if k in _RENDER_SECTIONS and isinstance(v, bool)}
        if cleaned:
            out["sections"] = cleaned
    order = raw.get("order")
    if isinstance(order, list):
        cleaned_order: list[str] = []
        for key in order:
            if isinstance(key, str) and key in _ORDERABLE_SECTIONS and key not in cleaned_order:
                cleaned_order.append(key)
        if cleaned_order:
            out["order"] = cleaned_order
    return out


# ── Classe principale ─────────────────────────────────────────────────────────

class DocxReport:
    def __init__(
        self,
        ctx: dict,
        participants: list,
        speaker_stats: dict,
        quality: dict,
        srt_text: str,
        structured_data: dict | None = None,
        render_options: dict | None = None,
        logo_bytes: bytes | None = None,
        summary_stale: bool = False,
    ):
        self.logo_bytes = logo_bytes
        # §5.2 : verbatim édité APRÈS la génération de la synthèse (marqueur posé par
        # l'éditeur SRT) → mention honnête dans le document, levée à la resync LLM.
        self.summary_stale = bool(summary_stale)
        self.ctx = ctx
        # Langue des livrables (Axe B) : pilote les libellés de chrome du DOCX.
        self.language: str = (ctx or {}).get("language") or "fr"
        self.L: dict[str, str] = _docx_labels(self.language)
        self.participants: list[dict] = participants if isinstance(participants, list) else []
        self.speakers: list[dict] = (speaker_stats or {}).get("speakers", [])
        self.quality = quality or {}
        self.srt_entries = _parse_srt(srt_text)
        self.duration_s = _srt_duration_seconds(srt_text)
        self.merged = self._merge_participants()
        self.structured_data: dict = structured_data if isinstance(structured_data, dict) else {}
        self.meeting_type: str = ctx.get("meeting_type", "") if ctx else ""
        self.type_specific_data: dict = ctx.get("type_specific_data") or {}
        # Fiche MATÉRIALISÉE d'un type personnalisé (posée par l'étape 4 dans
        # meeting_context["custom_type"]) : le rendu ne résout jamais un template en
        # base — la fiche du job fait foi, même si le template a été supprimé.
        raw_custom = ctx.get("custom_type") if ctx else None
        self.custom_type: dict = raw_custom if isinstance(raw_custom, dict) else {}
        custom_behavior = self.custom_type.get("behavior") or {}
        self.has_quorum: bool = self.meeting_type in _CSE_TYPES or bool(custom_behavior.get("quorum"))
        self.auto_confidential: bool = (
            self.meeting_type in _AUTO_CONFIDENTIEL or bool(custom_behavior.get("confidential"))
        )
        self.render_options: dict = _sanitize_render_options(render_options)
        theme_key = self.render_options.get("theme")
        if theme_key:
            self.theme: _DocxTheme = _THEMES[theme_key]
        elif self.custom_type.get("palette"):
            self.theme = _theme_from_definition(self.custom_type)
        else:
            self.theme = _get_theme(self.meeting_type)
        # Auto-confidentialité pour certains types
        if self.auto_confidential and not ctx.get("sensitivity"):
            self.ctx = dict(ctx)
            self.ctx["sensitivity"] = "high"

    # ── Fusion participants ───────────────────────────────────────────────────

    def _merge_participants(self) -> list[dict]:
        spk_map = {s["mapped_to"]: s for s in self.speakers if s.get("mapped_to")}
        total_time = sum(s.get("speaking_time_seconds", 0.0) for s in self.speakers)
        result: list[dict] = []

        for p in self.participants:
            spk = spk_map.get(p.get("id"), {})
            time_s = float(spk.get("speaking_time_seconds", 0))
            pct = round(100 * time_s / max(total_time, 0.001))
            result.append({
                "name": (p.get("name") or spk.get("mapped_name") or "—").strip(),
                "function": p.get("function", ""),
                "service": p.get("service", ""),
                "role": p.get("role", ""),
                "is_animator": bool(p.get("is_animator", False)),
                "time_s": time_s,
                "time_pct": pct,
                "turns": spk.get("turn_count", "—"),
            })

        mapped_ids = {p.get("id") for p in self.participants}
        for spk in self.speakers:
            if spk.get("mapped_to") not in mapped_ids:
                time_s = float(spk.get("speaking_time_seconds", 0))
                pct = round(100 * time_s / max(total_time, 0.001))
                result.append({
                    "name": (spk.get("mapped_name") or spk.get("speaker_id") or "—"),
                    "function": "",
                    "service": "",
                    "role": "",
                    "is_animator": False,
                    "time_s": time_s,
                    "time_pct": pct,
                    "turns": spk.get("turn_count", "—"),
                })
        return result

    # ── Build ─────────────────────────────────────────────────────────────────

    def _section_enabled(self, key: str) -> bool:
        # Surcharge par job (render_options) > défauts de la fiche du type > actif.
        job_sections = self.render_options.get("sections", {})
        if key in job_sections:
            return bool(job_sections[key])
        type_sections = (self.custom_type.get("sections") or {}).get("enabled") or {}
        if key in type_sections:
            return bool(type_sections[key])
        return True

    def _resolve_section_order(self) -> list[str]:
        """Ordre effectif des unités : job (render_options.order) > fiche du type >
        défaut historique. « contexte » et « pv » sont réinjectés s'ils manquent —
        déplaçables, jamais supprimables ; les unités par défaut omises sont
        rendues À LA SUITE (dans l'ordre historique)."""
        raw = self.render_options.get("order")
        if raw is None:
            raw = (self.custom_type.get("sections") or {}).get("order")
        if not isinstance(raw, list) or not raw:
            return list(_SECTION_ORDER_DEFAULT)
        order: list[str] = []
        for key in raw:
            if isinstance(key, str) and key in _ORDERABLE_SECTIONS and key not in order:
                order.append(key)
        for key in _SECTION_ORDER_DEFAULT:
            if key not in order:
                order.append(key)
        return order

    def build(self) -> DocumentT:
        doc = Document()
        self._setup_document(doc)
        self._cover_page(doc)
        self._page_break(doc)
        order = self._resolve_section_order()
        standalone = {"synthese", "champs_type"} & set(order)
        num = 1
        for key in order:
            if key == "contexte":
                self._section_context(
                    doc, number=f"{num}.",
                    include_synthese="synthese" not in standalone,
                    include_champs="champs_type" not in standalone,
                )
                num += 1
            elif key == "synthese":
                if self._section_synthese(doc, number=f"{num}."):
                    num += 1
            elif key == "champs_type":
                if self._section_champs_type(doc, number=f"{num}."):
                    num += 1
            elif key == "pv":
                num += self._section_enriched(doc, start=num)
            elif self._section_enabled(key):
                render = {"participants": self._section_participants,
                          "transcript": self._section_transcript,
                          "quality": self._section_quality}[key]
                render(doc, base=num)
                num += 1
        self._setup_footer(doc)
        return doc

    # ── Setup ─────────────────────────────────────────────────────────────────

    def _setup_document(self, doc: DocumentT) -> None:
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

    # ── Page de garde ─────────────────────────────────────────────────────────

    def _cover_page(self, doc: DocumentT) -> None:  # noqa: C901
        ctx    = self.ctx
        theme  = self.theme
        title  = ctx.get("title") or "Sans titre"
        mtype  = ctx.get("meeting_type", "")
        date   = _fmt_date(ctx.get("date", ""))
        svc    = ctx.get("service", "") or ""
        lang   = _LANG_LABELS.get(ctx.get("language", "fr"), ctx.get("language", ""))
        sensitivity = ctx.get("sensitivity", "normal")
        score  = self.quality.get("quality_score")
        ts     = self.type_specific_data

        # ── 0. Logo de l'installation/du type (fiche personnalisée, lot C) ────
        if self.logo_bytes:
            try:
                import io
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_logo.add_run().add_picture(io.BytesIO(self.logo_bytes), height=Cm(1.8))
                p_logo.paragraph_format.space_after = Pt(6)
            except Exception:
                # Logo illisible (fichier altéré) : la couverture se rend sans lui —
                # un livrable final ne plante jamais pour un élément décoratif.
                pass
        # ── 1. Bandeau principal (couleur signature du type) ─────────────────
        hdr = doc.add_table(rows=1, cols=1)
        _table_full_width(hdr)
        _table_no_borders(hdr)
        hdr_cell = hdr.cell(0, 0)
        _cell_bg(hdr_cell, theme.primary)
        _cell_margins(hdr_cell, top=260, bottom=260, left=360, right=360)
        p_hdr = hdr_cell.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Bandeau par défaut localisé ; un bandeau de type personnalisé (authoré) est conservé.
        # Bandeau : défaut localisé via la table ; bandeau typé localisé via le catalogue
        # (repli sur la forme française authorée pour un type personnalisé).
        if theme.banner_text == _DOCX_LABELS["fr"]["banner"]:
            banner_text = self.L["banner"]
        else:
            banner_text = localized_type_display(mtype, self.language, "banner_text", theme.banner_text)
        r_hdr = p_hdr.add_run(banner_text)
        r_hdr.font.color.rgb = _WHITE
        r_hdr.font.bold = True
        r_hdr.font.size = Pt(12)
        r_hdr.font.name = "Calibri"

        # ── 2. Trait d'accent mince sous le bandeau ──────────────────────────
        acc = doc.add_table(rows=1, cols=1)
        _table_full_width(acc)
        _table_no_borders(acc)
        acc_cell = acc.cell(0, 0)
        _cell_bg(acc_cell, theme.accent)
        _cell_margins(acc_cell, top=24, bottom=24, left=0, right=0)
        acc_cell.paragraphs[0].add_run("")

        # ── 3. Badge CONFIDENTIEL / CRISE (si applicable) ────────────────────
        is_confidentiel = (sensitivity == "high") or self.auto_confidential
        is_crise = mtype == "Réunion de crise"
        if is_confidentiel or is_crise:
            badge_color = _RED if is_crise else RGBColor(0x6A, 0x1B, 0x9A)
            badge_text  = self.L["badge_crise"] if is_crise else self.L["badge_confidentiel"]
            conf_t = doc.add_table(rows=1, cols=1)
            _table_full_width(conf_t)
            _table_no_borders(conf_t)
            conf_cell = conf_t.cell(0, 0)
            _cell_bg(conf_cell, badge_color)
            _cell_margins(conf_cell, top=80, bottom=80, left=200, right=200)
            p_conf = conf_cell.paragraphs[0]
            p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_conf = p_conf.add_run(badge_text)
            r_conf.font.color.rgb = _WHITE
            r_conf.font.bold = True
            r_conf.font.size = Pt(9)
            r_conf.font.name = "Calibri"

        # ── 4. Titre principal ───────────────────────────────────────────────
        doc.add_paragraph()
        doc.add_paragraph()
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_title = p_title.add_run(title.upper())
        run_title.font.size = Pt(22)
        run_title.font.bold = True
        run_title.font.color.rgb = theme.primary
        run_title.font.name = "Calibri"
        _para_bottom_border(p_title, theme.accent, sz=10)

        # ── 5. Sous-titre contextuel (nom projet, objet CSE extra…) ──────────
        subtitle_parts: list[str] = []
        if mtype in ("Point projet", "Réunion projet") and ts.get("nom_projet"):
            subtitle_parts.append(ts["nom_projet"])
            if ts.get("phase_jalon"):
                subtitle_parts.append(ts["phase_jalon"])
        elif self.has_quorum and ts.get("objet_seance"):
            subtitle_parts.append(ts["objet_seance"])
        elif mtype == "Réunion client" and ts.get("nom_client"):
            subtitle_parts.append(ts["nom_client"])
        elif mtype == "Réunion de crise" and ts.get("nature_incident"):
            subtitle_parts.append(ts["nature_incident"])
        elif mtype == "Entretien individuel" and ts.get("poste_evalue"):
            subtitle_parts.append(f"Entretien — {ts['poste_evalue']}")

        if subtitle_parts:
            p_sub = doc.add_paragraph()
            r_sub = p_sub.add_run("  ".join(subtitle_parts))
            r_sub.font.size = Pt(12)
            r_sub.font.color.rgb = theme.accent
            r_sub.font.italic = True
            r_sub.font.name = "Calibri"
            p_sub.paragraph_format.space_before = Pt(4)
            p_sub.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()

        # ── 6. Métadonnées en 2 colonnes soignées ────────────────────────────
        meta_rows: list[tuple[str, str]] = []
        if date and date != "—":
            meta_rows.append((self.L["date"], date))
        if self.duration_s:
            meta_rows.append((self.L["duration"], _fmt_duration(self.duration_s)))
        if mtype:
            meta_rows.append((self.L["type"], localized_type_display(mtype, self.language, "name", mtype)))
        if svc:
            meta_rows.append((self.L["service"], svc))
        if lang:
            meta_rows.append((self.L["language"], lang))
        # Champs type-spécifiques clés sur la couverture
        if self.has_quorum:
            if ts.get("president_seance"):
                meta_rows.append(("Président de séance", ts["president_seance"]))
            if ts.get("secretaire_seance"):
                meta_rows.append(("Secrétaire de séance", ts["secretaire_seance"]))
            if ts.get("ref_pv_precedent"):
                meta_rows.append(("Réf. PV précédent", ts["ref_pv_precedent"]))
        elif mtype in ("Point projet", "Réunion projet"):
            if ts.get("chef_de_projet"):
                meta_rows.append(("Chef de projet", ts["chef_de_projet"]))
            if ts.get("sprint"):
                meta_rows.append(("Sprint", ts["sprint"]))
        elif mtype == "CODIR / COMEX":
            pass  # ordre du jour dans le document
        elif mtype == "Réunion client":
            if ts.get("ref_contrat"):
                meta_rows.append(("Réf. contrat", ts["ref_contrat"]))
        elif mtype == "Entretien individuel":
            if ts.get("periode_evaluee"):
                meta_rows.append(("Période", ts["periode_evaluee"]))
            if ts.get("evaluateur"):
                meta_rows.append(("Évaluateur", ts["evaluateur"]))

        if meta_rows:
            meta_t = doc.add_table(rows=len(meta_rows), cols=2)
            _table_full_width(meta_t)
            _table_no_borders(meta_t)
            for i, (lbl, val) in enumerate(meta_rows):
                cells = meta_t.rows[i].cells
                _cell_margins(cells[0], top=36, bottom=36, left=0, right=80)
                _cell_margins(cells[1], top=36, bottom=36, left=80, right=0)
                r_l = cells[0].paragraphs[0].add_run(lbl)
                r_l.font.size = Pt(9.5)
                r_l.font.bold = True
                r_l.font.color.rgb = _GREY_DARK
                r_l.font.name = "Calibri"
                r_v = cells[1].paragraphs[0].add_run(val)
                r_v.font.size = Pt(9.5)
                r_v.font.color.rgb = theme.primary
                r_v.font.name = "Calibri"

        # ── 7. Quorum CSE (encadré visuel fort) ──────────────────────────────
        if self.has_quorum:
            try:
                presents = int(ts.get("membres_presents") or 0)
                total    = int(ts.get("membres_total") or 0)
                if presents and total:
                    quorum_ok  = presents > total / 2
                    pct        = round(100 * presents / total)
                    quorum_txt = (f"✓  Quorum atteint — {presents}/{total} membres présents ({pct}%)"
                                  if quorum_ok
                                  else f"✗  Quorum NON atteint — {presents}/{total} membres présents ({pct}%)")
                    doc.add_paragraph()
                    q_t = doc.add_table(rows=1, cols=1)
                    _table_full_width(q_t)
                    _table_no_borders(q_t)
                    q_cell = q_t.cell(0, 0)
                    _cell_bg(q_cell, _GREEN if quorum_ok else RGBColor(0xFF, 0xEB, 0xEE))
                    _cell_margins(q_cell, top=120, bottom=120, left=200, right=200)
                    p_q = q_cell.paragraphs[0]
                    p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_q = p_q.add_run(quorum_txt)
                    r_q.font.size = Pt(10)
                    r_q.font.bold = True
                    r_q.font.color.rgb = _WHITE if quorum_ok else _RED
                    r_q.font.name = "Calibri"
            except (ValueError, TypeError):
                pass

        # ── 8. Pied de couverture ─────────────────────────────────────────────
        doc.add_paragraph()
        doc.add_paragraph()
        p_gen = doc.add_paragraph()
        p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_gen = p_gen.add_run(f"{self.L['generated_by']}  ▪  {datetime.today().strftime('%d/%m/%Y')}")
        r_gen.font.size = Pt(8.5)
        r_gen.font.color.rgb = _GREY_DARK
        r_gen.font.name = "Calibri"

        if score is not None:
            score_color = _GREEN if score >= 85 else _ORANGE if score >= 65 else _RED
            r_score = p_gen.add_run(f"  ▪  {self.L['quality']} : {score}/100")
            r_score.font.size = Pt(8.5)
            r_score.font.color.rgb = score_color
            r_score.font.bold = True
            r_score.font.name = "Calibri"

        if theme.cover_badge:
            p_badge = doc.add_paragraph()
            p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _badge = localized_type_display(mtype, self.language, "badge", theme.cover_badge)
            r_badge = p_badge.add_run(f"[ {_badge} ]")
            r_badge.font.size = Pt(8)
            r_badge.font.color.rgb = theme.accent
            r_badge.font.bold = True
            r_badge.font.name = "Calibri"
            p_badge.paragraph_format.space_before = Pt(2)

    # ── Helpers section ───────────────────────────────────────────────────────

    def _section_heading(self, doc: DocumentT, number: str, label: str) -> None:
        theme = self.theme
        doc.add_paragraph()
        p = doc.add_paragraph()
        # Numéro en couleur accent
        r_num = p.add_run(f"{number}  ")
        r_num.font.size = Pt(13)
        r_num.font.bold = True
        r_num.font.color.rgb = theme.accent
        r_num.font.name = "Calibri"
        # Libellé en couleur primaire
        r_lbl = p.add_run(label.upper())
        r_lbl.font.size = Pt(13)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = theme.primary
        r_lbl.font.name = "Calibri"
        # Trait bas coloré
        _para_bottom_border(p, theme.accent, sz=6)

    @staticmethod
    def _meta_row(doc: DocumentT, label: str, value: str) -> None:
        p = doc.add_paragraph()
        r_lbl = p.add_run(f"{label} : ")
        r_lbl.font.size = Pt(10)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = _GREY_DARK
        r_lbl.font.name = "Calibri"
        r_val = p.add_run(value)
        r_val.font.size = Pt(10)
        r_val.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _page_break(doc: DocumentT) -> None:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── Section 1 : Contexte ──────────────────────────────────────────────────

    def _summary_text(self) -> str:
        # Priorité : édition manuelle (étape 4) > synthèse harmonisée sur le glossaire
        # validé (post-correction) > synthèse brute de la LLM (pré-correction).
        text = (
            self.ctx.get("summary")
            or self.ctx.get("summary_harmonized")
            or self.ctx.get("summary_llm")
            or ""
        ).strip()
        # Rédigée AVANT la validation des locuteurs (autostart) : les jetons
        # SPEAKER_XX sont remplacés au RENDU par les noms validés (artefact intact).
        # Différé : arête MONTANTE exports→workflow (cycle d'__init__ via phases/export).
        from transcria.workflow.speaker_projection import substitute_speaker_names

        return substitute_speaker_names(text, self.ctx.get("speaker_mapping"))

    def _section_context(self, doc: DocumentT, number: str = "1.",
                         include_synthese: bool = True, include_champs: bool = True) -> None:
        ctx = self.ctx
        self._section_heading(doc, number, self.L["sec_context"])

        # `or ""` AVANT .strip() : une clé présente à `null` renvoie None (le défaut de
        # .get ne couvre que les clés ABSENTES) → None.strip() planterait tout le rapport.
        # `str(... or "")` couvre null / absent / vide / non-chaîne (un livrable final ne
        # plante jamais sur un champ de contexte mal typé).
        topic = str(ctx.get("topic") or "").strip()
        objective = str(ctx.get("objective") or "").strip()
        notes = str(ctx.get("notes") or "").strip()
        summary = self._summary_text() if include_synthese else ""

        if topic:
            self._meta_row(doc, self.L["topic"], topic)
        if objective:
            self._meta_row(doc, self.L["objective"], objective)
        if notes and notes.lower() not in ("n/a", "n/a — scène unique de dialogue.", ""):
            self._meta_row(doc, self.L["notes"], notes)

        if summary:
            doc.add_paragraph()
            p_head = doc.add_paragraph()
            r = p_head.add_run(self.L["synthese"])
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = self.theme.primary
            r.font.name = "Calibri"
            p_head.paragraph_format.space_after = Pt(4)
            self._render_synthese_body(doc, summary)

        if include_champs:
            # Champs utilisateur spécifiques au type (président CSE, nom projet, etc.)
            self._section_type_specific(doc)

    def _render_synthese_body(self, doc: DocumentT, summary: str) -> None:
        if self.summary_stale:
            note = doc.add_paragraph()
            run = note.add_run(self.L["summary_stale_note"])
            run.italic = True
            run.font.size = Pt(9)
        # Extraire juste le paragraphe "Synthèse" si présent dans un markdown
        synth = _extract_synthese(summary, self.language)
        for raw in synth.splitlines():
            line = raw.strip()
            if not line:
                continue
            # Intertitre markdown (## …) → ligne en gras, légèrement détachée.
            heading = re.match(r"^#{1,6}\s+(.*)$", line)
            if heading:
                p = doc.add_paragraph()
                _add_markdown_runs(p, heading.group(1).strip(), size=10.5, bold_all=True)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                continue
            # Puce markdown (- … ou * …).
            bullet = re.match(r"^[-*]\s+(.*)$", line)
            p = doc.add_paragraph()
            if bullet:
                _add_markdown_runs(p, "•  " + bullet.group(1).strip(), size=10)
                p.paragraph_format.left_indent = Cm(0.5)
            else:
                # Paragraphe de prose : le gras **…** (intertitres en début de
                # paragraphe) est rendu en gras réel au lieu d'être supprimé.
                _add_markdown_runs(p, line, size=10)
            p.paragraph_format.space_after = Pt(4)

    def _section_synthese(self, doc: DocumentT, number: str) -> bool:
        """Synthèse en SECTION AUTONOME (« résumé exécutif en premier ») — rendue
        seulement si un ordre personnalisé la cite ET qu'une synthèse existe."""
        summary = self._summary_text()
        if not summary:
            return False
        self._section_heading(doc, number, self.L["synthese"])
        self._render_synthese_body(doc, summary)
        return True

    def _section_champs_type(self, doc: DocumentT, number: str) -> bool:
        """Champs du type en section autonome (ordre personnalisé)."""
        if not self._has_type_specific_rows():
            return False
        self._section_heading(doc, number, self.L["sec_specific"])
        self._section_type_specific(doc)
        return True

    def _has_type_specific_rows(self) -> bool:
        ts = self.type_specific_data
        return bool(ts) and any(v is not None and str(v).strip() for v in ts.values())

    # ── Section 1b : Données type-spécifiques (champs utilisateur) ──────────────

    def _section_type_specific(self, doc: DocumentT) -> None:
        """Affiche les champs saisis par l'utilisateur pour ce type de réunion.

        Absent si aucun champ n'a été rempli.
        Pour CSE : indicateur de quorum calculé automatiquement.
        """
        ts = self.type_specific_data
        if not ts:
            return

        # Filtrer les champs non vides
        non_empty = {k: v for k, v in ts.items() if v is not None and str(v).strip()}
        if not non_empty:
            return

        # Libellés courts par clé — depuis le catalogue (`short_label` sinon `label`),
        # complétés par la fiche du type personnalisé du job le cas échéant ; une clé
        # inconnue (donnée ancienne, type supprimé) garde le repli générique.
        labels_by_key = dict(localized_field_labels(self.language))
        for field in self.custom_type.get("fields") or []:
            if isinstance(field, dict) and field.get("key"):
                labels_by_key[field["key"]] = str(field.get("short_label") or field.get("label") or field["key"])

        doc.add_paragraph()
        # Tableau compact sans bordures extérieures
        rows_data: list[tuple[str, str]] = []

        for key, val in non_empty.items():
            label = labels_by_key.get(key, key.replace("_", " ").capitalize())
            # Ordre du jour : chaque ligne → item
            if key == "ordre_du_jour_items":
                for i, line in enumerate(str(val).splitlines()):
                    line = line.strip()
                    if line:
                        rows_data.append((f"{self.L['odj_prefix']} {i+1}" if i == 0 else "", line))
                continue
            rows_data.append((label, str(val).strip()))

        # Quorum calculé (types à drapeau behavior.quorum)
        if self.has_quorum:
            try:
                presents = int(non_empty.get("membres_presents", 0))
                total    = int(non_empty.get("membres_total", 0))
                if presents and total:
                    pct    = round(100 * presents / total)
                    quorum = self.L["quorum_reached"] if presents > total / 2 else self.L["quorum_not_reached"]
                    rows_data.append((self.L["quorum"], f"{quorum} ({presents}/{total} — {pct}%)"))
            except (ValueError, TypeError):
                pass

        if not rows_data:
            return

        table = doc.add_table(rows=len(rows_data), cols=2)
        _table_full_width(table)
        _table_no_borders(table)

        for i, (label, val) in enumerate(rows_data):
            cells = table.rows[i].cells
            _cell_margins(cells[0], top=30, bottom=30, left=0, right=60)
            _cell_margins(cells[1], top=30, bottom=30, left=60, right=0)

            r_lbl = cells[0].paragraphs[0].add_run(label)
            r_lbl.font.size = Pt(9.5)
            r_lbl.font.bold = True
            r_lbl.font.color.rgb = _GREY_DARK
            r_lbl.font.name = "Calibri"

            # Détection de couleur indépendante de la langue (symboles ✓/✗ des libellés quorum).
            color = _GREEN if "✓" in val else _RED if "✗" in val else self.theme.primary
            r_val = cells[1].paragraphs[0].add_run(val)
            r_val.font.size = Pt(9.5)
            r_val.font.color.rgb = color
            r_val.font.name = "Calibri"

    # ── Section 1c : Données enrichies LLM (décisions, actions, votes…) ─────────

    def _section_enriched(self, doc: DocumentT, start: int = 2) -> int:
        """Sections issues de l'extraction LLM structurée.

        Principe : **une donnée extraite n'est jamais cachée**. Toute section
        s'affiche dès qu'elle contient des éléments, quel que soit le type de
        réunion. Le type pilote uniquement le thème visuel et les champs de
        saisie (`type_specific_data`), pas la rétention du contenu extrait.

        Ordre fixe inspiré d'un procès-verbal : agenda → décisions → votes →
        résolutions → actions → blocages → reports. L'absence totale est
        silencieuse (aucun placeholder).
        """
        sd = self.structured_data
        section_num = start  # numéro de la première section PV rendue

        def _as_str_items(val: Any) -> list[str]:
            # Défense : la structure canonique est une liste de chaînes, mais le JSON
            # relu par la relecture finale peut dévier (dicts, scalaires) — un rapport
            # final ne plante jamais pour autant, on coerce en texte.
            if isinstance(val, str):
                return [val.strip()] if val.strip() else []
            if isinstance(val, (list, tuple)):
                # Textes rédigés par la LLM (souvent AVANT la validation des
                # locuteurs) : mêmes substitutions de noms que la synthèse.
                from transcria.workflow.speaker_projection import substitute_speaker_names

                mapping = self.ctx.get("speaker_mapping")
                return [substitute_speaker_names(s, mapping)
                        for s in (str(item).strip() for item in val) if s]
            return []

        # (label, items) dans l'ordre PV — chaque section affichée si non vide
        ordered = [
            (self.L["sd_agenda"],      _as_str_items(sd.get("points_odj"))),
            (self.L["sd_decisions"],   _as_str_items(sd.get("decisions"))),
            (self.L["sd_votes"],       _as_str_items(sd.get("votes"))),
            (self.L["sd_resolutions"], _as_str_items(sd.get("resolutions"))),
            (self.L["sd_actions"],     _as_str_items(sd.get("actions"))),
            (self.L["sd_blockers"],    _as_str_items(sd.get("blocages"))),
            (self.L["sd_deferred"],    _as_str_items(sd.get("reports"))),
        ]
        # Champs d'extraction du type personnalisé (fiche matérialisée, lot D) —
        # rendus APRÈS les blocs PV universels, même règle « non vide ⇒ affiché ».
        for field in self.custom_type.get("extract_fields") or []:
            if isinstance(field, dict) and field.get("key"):
                label = str(field.get("label") or field["key"])
                ordered.append((label, _as_str_items(sd.get(field["key"]))))
        shown: list[tuple[str, list[str]]] = [
            (label, items) for label, items in ordered if items
        ]

        if not shown:
            return 0

        for label, items in shown:
            self._section_heading(doc, f"{section_num}.", label)
            section_num += 1
            for item in items:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)
                run_bullet = p.add_run("▸  ")
                run_bullet.font.color.rgb = self.theme.accent
                run_bullet.font.size = Pt(9)
                # Gras markdown de la LLM rendu en vrai gras (pas de ** littéraux).
                _add_markdown_runs(p, item, size=10)

        # Prochaine date — footer discret si mentionnée
        if sd.get("prochaine_date"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            r_lbl = p.add_run(self.L["next_meeting"])
            r_lbl.font.size = Pt(9)
            r_lbl.font.bold = True
            r_lbl.font.color.rgb = _GREY_DARK
            r_lbl.font.name = "Calibri"
            r_val = p.add_run(str(sd["prochaine_date"]))
            r_val.font.size = Pt(9)
            r_val.font.color.rgb = self.theme.accent
            r_val.font.name = "Calibri"

        return len(shown)

    # ── Section N : Participants ──────────────────────────────────────────────

    def _section_participants(self, doc: DocumentT, base: int = 2) -> None:
        self._section_heading(doc, f"{base}.", self.L["sec_participants"])

        if not self.merged:
            doc.add_paragraph(self.L["no_participants"])
            return

        has_function = any(p["function"] for p in self.merged)
        has_service  = any(p["service"]  for p in self.merged)
        has_animator = any(p["is_animator"] for p in self.merged)

        headers = [self.L["th_name"]]
        if has_function:
            headers.append(self.L["th_function"])
        if has_service:
            headers.append(self.L["th_service"])
        headers += [self.L["th_role"], self.L["th_speaking_time"], self.L["th_turns"]]
        if has_animator:
            headers.append("")

        n_cols = len(headers)
        table = doc.add_table(rows=1 + len(self.merged), cols=n_cols)
        _table_full_width(table)
        _table_thin_borders(table)

        # ── En-tête ──────────────────────────────────────────────────────────
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            _cell_bg(hdr_cells[i], self.theme.primary)
            _cell_margins(hdr_cells[i], top=80, bottom=80, left=100, right=100)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.bold = True
            run.font.color.rgb = _WHITE
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

        # ── Lignes données ────────────────────────────────────────────────────
        for row_i, participant in enumerate(self.merged):
            row = table.rows[row_i + 1]
            bg = self.theme.light if row_i % 2 == 0 else _WHITE

            data: list[tuple[str, bool, RGBColor]] = []  # (text, bold, color)

            data.append((participant["name"], True, self.theme.primary))
            if has_function:
                data.append((participant["function"] or "—", False, _GREY_DARK))
            if has_service:
                data.append((participant["service"] or "—", False, _GREY_DARK))
            role = participant["role"][:80] + "…" if len(participant["role"]) > 80 else participant["role"]
            data.append((role or "—", False, _GREY_DARK))
            time_label = (
                f"{_fmt_time(participant['time_s'])} ({participant['time_pct']}%)"
                if participant["time_s"] > 0 else "—"
            )
            data.append((time_label, False, _GREY_DARK))
            data.append((str(participant["turns"]), False, _GREY_DARK))
            if has_animator:
                data.append(("★ Animateur" if participant["is_animator"] else "", True, self.theme.accent))

            for col, (text, bold, color) in enumerate(data):
                cell = row.cells[col]
                _cell_bg(cell, bg)
                _cell_margins(cell, top=60, bottom=60, left=100, right=100)
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(9.5)
                run.font.bold = bold
                run.font.color.rgb = color
                run.font.name = "Calibri"

    # ── Section 3 : Transcription ─────────────────────────────────────────────

    def _section_transcript(self, doc: DocumentT, base: int = 3) -> None:
        self._section_heading(doc, f"{base}.", self.L["sec_transcription"])

        if not self.srt_entries:
            doc.add_paragraph(self.L["no_transcription"])
            return

        table = doc.add_table(rows=len(self.srt_entries), cols=3)
        _table_full_width(table)
        _table_no_borders(table)

        for i, entry in enumerate(self.srt_entries):
            row = table.rows[i]
            bg = _GREY_LIGHT if i % 2 == 0 else _WHITE

            # Col 0 — timestamp
            ts_cell = row.cells[0]
            _cell_bg(ts_cell, bg)
            _cell_margins(ts_cell, top=40, bottom=40, left=0, right=60)
            p0 = ts_cell.paragraphs[0]
            r0 = p0.add_run(entry["timestamp"])
            r0.font.size = Pt(8)
            r0.font.italic = True
            r0.font.color.rgb = _GREY_DARK
            r0.font.name = "Consolas"

            # Col 1 — locuteur
            spk_cell = row.cells[1]
            _cell_bg(spk_cell, bg)
            _cell_margins(spk_cell, top=40, bottom=40, left=60, right=80)
            p1 = spk_cell.paragraphs[0]
            r1 = p1.add_run(entry["speaker"])
            r1.font.size = Pt(9)
            r1.font.bold = True
            r1.font.color.rgb = self.theme.accent if entry["speaker"] else _GREY_DARK
            r1.font.name = "Calibri"

            # Col 2 — texte
            txt_cell = row.cells[2]
            _cell_bg(txt_cell, bg)
            _cell_margins(txt_cell, top=40, bottom=40, left=80, right=0)
            p2 = txt_cell.paragraphs[0]
            r2 = p2.add_run(entry["text"])
            r2.font.size = Pt(9.5)
            r2.font.name = "Calibri"

    # ── Section 4 : Points à vérifier (conditionnelle) ────────────────────────

    # Libellés français des checks sans rendu dédié (repli générique — un avertissement
    # du rapport qualité n'est JAMAIS caché au lecteur du document final).
    # Libellés génériques des contrôles qualité → désormais dans _DOCX_LABELS (clés « chk_* »),
    # résolus par langue via self.L (Axe B).

    def _section_quality(self, doc: DocumentT, base: int = 4) -> None:
        checks = self.quality.get("checks", [])
        points: list[tuple[str, str]] = []  # (emoji_label, description)

        for check in checks:
            ctype = check.get("type")
            sev   = check.get("severity", "info")
            if sev == "info":
                continue

            if ctype == "low_coverage":
                ratio = check.get("ratio", 1.0)
                if ratio < 0.85:
                    pct = round(ratio * 100)
                    points.append((self.L["q_coverage"], self.L["d_coverage"].format(pct=pct)))

            elif ctype == "audio_problem_segments":
                examples = check.get("examples", [])
                for ex in examples:
                    label = ex.get("label", self.L["q_anomaly"])
                    s = ex.get("start_label", "")
                    e = ex.get("end_label", "")
                    points.append((self.L["q_relisten"], f"{s} → {e} ({label})"))

            elif ctype == "unresolved_lexicon_variants":
                for ev in check.get("exact_variants", []):
                    points.append((self.L["q_term_validate"],
                                   self.L["d_term_validate"].format(term=ev['term'], variant=ev['variant'])))
                for cf in check.get("close_forms", []):
                    points.append((self.L["q_spelling"],
                                   self.L["d_spelling"].format(form=cf['form'], term=cf['term'])))

            elif ctype == "speaker_name_violations":
                # severity=error : un nom de locuteur a été ALTÉRÉ dans le SRT corrigé —
                # information capitale pour le relecteur du document final.
                for v in check.get("violations", [])[:10]:
                    points.append((
                        self.L["q_altered_name"],
                        self.L["d_altered_name"].format(
                            sid=v.get('speaker_id', '?'), found=v.get('found', ''), expected=v.get('expected', '')),
                    ))

            elif ctype == "missing_lexicon_terms":
                for term in check.get("terms", [])[:10]:
                    points.append((self.L["q_missing_term"], self.L["d_missing_term"].format(term=term)))

            elif ctype == "unmapped_speakers":
                points.append((self.L["q_unmapped"], self.L["d_unmapped"].format(count=check.get('count', 0))))

            elif ctype == "foreign_segments":
                points.append((self.L["q_foreign"], self.L["d_foreign"].format(count=check.get('count', 0))))

            elif ctype == "non_latin_segments":
                points.append((self.L["q_non_latin"], self.L["d_non_latin"].format(count=check.get('count', 0))))

            else:
                label = self.L.get("chk_" + str(ctype), str(ctype).replace("_", " ").capitalize())
                count = check.get("count")
                desc = self.L["d_generic_count"].format(count=count) if count else self.L["d_generic_details"]
                points.append((f"⚠  {label}", desc))

        if not points:
            return

        self._section_heading(doc, f"{base}.", self.L["sec_review"])

        table = doc.add_table(rows=len(points), cols=2)
        _table_full_width(table)
        _table_no_borders(table)

        for i, (label, desc) in enumerate(points):
            cells = table.rows[i].cells
            _cell_bg(cells[0], _YELLOW_BG)
            _cell_bg(cells[1], _YELLOW_BG)
            _cell_margins(cells[0], top=60, bottom=60, left=100, right=80)
            _cell_margins(cells[1], top=60, bottom=60, left=80, right=100)

            r0 = cells[0].paragraphs[0].add_run(label)
            r0.font.size = Pt(9.5)
            r0.font.bold = True
            r0.font.color.rgb = RGBColor(0x7B, 0x36, 0x06)
            r0.font.name = "Calibri"

            r1 = cells[1].paragraphs[0].add_run(desc)
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = RGBColor(0x7B, 0x36, 0x06)
            r1.font.name = "Calibri"

    # ── Pied de page ──────────────────────────────────────────────────────────

    def _setup_footer(self, doc: DocumentT) -> None:
        theme  = self.theme
        title  = (self.ctx.get("title") or "TranscrIA")[:40]
        date   = _fmt_date(self.ctx.get("date", ""))
        score  = self.quality.get("quality_score")

        section = doc.sections[0]
        # Page de garde vierge (pas de « Page 1/N » sur la couverture) : le pied de
        # page ci-dessous ne s'applique qu'à partir de la page 2.
        section.different_first_page_header_footer = True
        footer  = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Pastille score qualité
        if score is not None:
            score_color = _GREEN if score >= 85 else _ORANGE if score >= 65 else _RED
            r_score = p.add_run(f"■ {score}/100   ")
            r_score.font.size = Pt(7.5)
            r_score.font.color.rgb = score_color
            r_score.font.name = "Calibri"

        # Trait couleur theme avant le titre
        r_accent = p.add_run("▪ ")
        r_accent.font.size = Pt(7.5)
        r_accent.font.color.rgb = theme.accent
        r_accent.font.name = "Calibri"

        footer_text = str((self.custom_type.get("branding") or {}).get("footer_text") or "").strip()
        prefix = f"{footer_text}  ·  " if footer_text else ""
        r_info = p.add_run(f"{prefix}TranscrIA  ·  {title}  ·  {date}  ·  Page ")
        r_info.font.size = Pt(7.5)
        r_info.font.color.rgb = _GREY_DARK
        r_info.font.name = "Calibri"

        r_pg = p.add_run()
        r_pg.font.size = Pt(7.5)
        r_pg.font.color.rgb = _GREY_DARK
        _add_page_number_field(r_pg)

        r_sep = p.add_run(" / ")
        r_sep.font.size = Pt(7.5)
        r_sep.font.color.rgb = _GREY_DARK
        r_sep.font.name = "Calibri"

        r_tot = p.add_run()
        r_tot.font.size = Pt(7.5)
        r_tot.font.color.rgb = _GREY_DARK
        _add_num_pages_field(r_tot)


# ── Extraction synthèse depuis markdown ──────────────────────────────────────

def _extract_synthese(text: str, language: str = "fr") -> str:
    """Extrait la section synthèse d'un markdown LLM (en-tête selon ``language``).

    Sans langue, cherche « ## Synthèse » (FR historique). En anglais, cherche
    « ## Summary » — sinon le repli déversait TOUT le summary.md (méta + termes +
    bloc JSON) dans le DOCX. Repli final : texte sans les titres markdown."""
    heading = re.escape(summary_markers(language)["summary_heading"].lstrip("# ").strip())
    m = re.search(rf"##\s*{heading}\s*\n(.*?)(?=\n##|\Z)", text, re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1).strip()
    # Repli FR historique conservé (résumés antérieurs), puis nettoyage des titres.
    m_fr = re.search(r"##\s*Synth[eè]se\s*\n(.*?)(?=\n##|\Z)", text, re.DOTALL | re.IGNORECASE)
    if m_fr:
        return m_fr.group(1).strip()
    cleaned = re.sub(r"^#{1,3}\s+.+$", "", text, flags=re.MULTILINE)
    return cleaned.strip()




def generate_docx_report(job_id: str, jobs_dir: str, output_path: Path) -> Path:
    """
    Génère le rapport DOCX pour un job terminé et l'écrit dans output_path.

    Retourne le chemin du fichier généré.
    """
    fs = JobFilesystem(jobs_dir, job_id)

    ctx           = fs.load_json("context/meeting_context.json") or {}
    participants  = fs.load_json("context/participants.json") or []
    speaker_stats = fs.load_json("speakers/speaker_stats.json") or {}
    speaker_map   = fs.load_json("speakers/speaker_mapping.json") or {}
    quality       = fs.load_json("quality/quality_report.json") or {}
    structured_data = ctx.get("structured_data") or {}

    srt_text = fs.load_text("metadata/transcription_corrigee.srt") or ""
    if not srt_text:
        srt_text = fs.load_text("metadata/transcription.srt") or ""

    try:
        render_options = fs.load_json("context/render_options.json") or {}
    except Exception:  # JSON corrompu = options par défaut, le rendu ne casse jamais
        render_options = {}

    # Logo du type personnalisé, matérialisé dans le job à l'étape 4 (comme la fiche).
    logo_bytes: bytes | None = None
    logo_path = fs.job_dir / "context" / "type_logo.png"
    if logo_path.is_file():
        try:
            logo_bytes = logo_path.read_bytes()
        except OSError:
            logo_bytes = None

    summary_stale = bool(fs.load_json("metadata/summary_stale.json") or {})
    # Mapping locuteurs validé → substitution des jetons SPEAKER_XX au rendu de la
    # synthèse (le ctx est le canal de données du rapport ; l'artefact reste intact).
    ctx["speaker_mapping"] = speaker_map
    report = DocxReport(ctx, participants, speaker_stats, quality, srt_text, structured_data,
                        render_options=render_options, logo_bytes=logo_bytes,
                        summary_stale=summary_stale)
    doc = report.build()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
